from __future__ import annotations

import asyncio
import sqlite3
from collections import Counter
from types import SimpleNamespace

from docx import Document

from app.external_graphrag import document_structure_counts
from app.legal_graphrag import LegalGraphBuilder
from app.services.retrieval import (
    RetrievalService,
    classify_retrieval_route,
    is_document_structure_query,
)


def test_builder_materializes_hierarchy_counts_and_inverse_edges(tmp_path) -> None:
    data_dir = tmp_path / "data"
    data_dir.mkdir()
    document = Document()
    document.add_paragraph("BỘ LUẬT")
    document.add_paragraph("LAO ĐỘNG")
    document.add_paragraph("Bộ luật số: 45/2019/QH14")
    document.add_paragraph("Chương I. QUY ĐỊNH CHUNG")
    document.add_paragraph("Điều 1. Phạm vi điều chỉnh")
    document.add_paragraph("1. Nội dung thứ nhất")
    document.add_paragraph("a) Điểm thứ nhất")
    document.add_paragraph("2. Nội dung thứ hai")
    document.add_paragraph("Điều 2. Đối tượng áp dụng")
    document.add_paragraph("1. Nội dung thứ ba")
    document.add_paragraph("a) Điểm thứ hai")
    document.save(data_dir / "Bộ-luật-45-2019-QH14.docx")

    builder = LegalGraphBuilder(data_dir, tmp_path / "storage")
    builder._parse_document(data_dir / "Bộ-luật-45-2019-QH14.docx")
    builder._finalize_node_text()
    builder._finalize_structure_metadata()
    builder._refresh_path_labels()
    builder._build_chunks()

    law = builder.nodes["doc:bo-luat-45-2019-qh14"]
    assert law["chapter_count"] == 1
    assert law["article_count"] == 2
    assert law["clause_count"] == 3
    assert law["point_count"] == 2
    assert law["first_article_number"] == "1"
    assert law["last_article_number"] == "2"

    relations = Counter(edge["relation"] for edge in builder.edges.values())
    assert relations["CÓ_CHƯƠNG"] == 1
    assert relations["CÓ_ĐIỀU"] == 2
    assert relations["CÓ_KHOẢN"] == 3
    assert relations["CÓ_ĐIỂM"] == 2

    structure = next(
        chunk
        for chunk in builder.chunks.values()
        if chunk["chunk_type"] == "document_structure"
    )
    assert document_structure_counts(structure) == {
        "chapters": 1,
        "sections": 0,
        "articles": 2,
        "clauses": 3,
        "points": 2,
        "first_article": "1",
        "last_article": "2",
    }

    builder.storage_dir.mkdir(parents=True, exist_ok=True)
    builder._write_sqlite()
    with sqlite3.connect(builder.db_path) as connection:
        stored = connection.execute(
            """
            SELECT article_count, clause_count, point_count
            FROM nodes
            WHERE node_id = 'doc:bo-luat-45-2019-qh14'
            """
        ).fetchone()
        assert stored == (2, 3, 2)


def test_structure_question_uses_graph_count_without_external_freshness() -> None:
    query = "Luật lao động 2019 có mấy điều"

    class _StructureStore:
        def document_structures(self, _: int) -> list[dict]:
            return [
                {
                    "chunk_id": "structure-labor-code",
                    "doc_id": "bo-luat-45-2019-qh14",
                    "node_id": "doc:bo-luat-45-2019-qh14",
                    "chunk_type": "document_structure",
                    "title": "Cấu trúc Bộ Luật Lao Động (45/2019/QH14)",
                    "citation": "Bộ Luật Lao Động (45/2019/QH14)",
                    "text": (
                        "Thống kê cấu trúc của Bộ Luật Lao Động. "
                        "STRUCTURE_COUNTS: chapters=17; sections=24; "
                        "articles=220; clauses=640; points=272; "
                        "first_article=1; last_article=220."
                    ),
                    "law_code": "45/2019/QH14",
                }
            ]

    service = RetrievalService(
        SimpleNamespace(retriever_backend="hybrid_rag")
    )
    service._store = _StructureStore()

    result = asyncio.run(service.lookup_document_structure(query))

    assert is_document_structure_query(query) is True
    assert classify_retrieval_route(query) == "multi_abstract"
    assert result is not None
    assert "**220 điều**" in result["answer"]
    assert "Điều 1 đến Điều 220" in result["answer"]
    assert "**640 khoản**" in result["answer"]
    assert "Graph" not in result["answer"]
    assert result["source"]["source_id"] == "S1"
    assert "document_structure_graph" in result["source"]["reasons"]
