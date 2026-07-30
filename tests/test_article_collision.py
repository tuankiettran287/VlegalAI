from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document

from app.legal_graphrag import LegalGraphBuilder


def create_mock_docx_with_duplicate_articles(target_path: Path) -> None:
    """Create a sample docx containing 'Điều 1' in main text AND 'Điều 1' in Phụ lục."""
    doc = Document()
    doc.add_paragraph("Nghị định 145/2020/NĐ-CP quy định chi tiết Bộ luật Lao động")
    doc.add_paragraph("Chương I")
    doc.add_paragraph("QUY ĐỊNH CHUNG")
    doc.add_paragraph("Điều 1. Phạm vi điều chỉnh")
    doc.add_paragraph("1. Nghị định này quy định chi tiết một số điều của Bộ luật Lao động.")

    doc.add_paragraph("PHỤ LỤC I")
    doc.add_paragraph("MẪU HỢP ĐỒNG LAO ĐỘNG")
    doc.add_paragraph("Điều 1. Công việc và địa điểm làm việc")
    doc.add_paragraph("1. Người lao động làm công việc chuyên môn tại văn phòng.")

    doc.save(str(target_path))


def test_article_node_id_collision_prevention():
    with tempfile.TemporaryDirectory() as tmp_dir:
        data_dir = Path(tmp_dir) / "data"
        storage_dir = Path(tmp_dir) / "storage"
        data_dir.mkdir()
        storage_dir.mkdir()

        mock_docx = data_dir / "nghi-dinh-145-2020-nd-cp.docx"
        create_mock_docx_with_duplicate_articles(mock_docx)

        builder = LegalGraphBuilder(data_dir, storage_dir)
        # Parse document without embedding
        builder.storage_dir.mkdir(parents=True, exist_ok=True)
        builder._parse_document(mock_docx)
        builder._finalize_node_text()
        builder._refresh_path_labels()

        # Find all nodes of type 'Điều'
        article_nodes = [node for node in builder.nodes.values() if node["node_type"] == "Điều"]

        # Must have at least 2 distinct Article nodes
        assert len(article_nodes) == 2, f"Expected 2 'Điều' nodes, got {len(article_nodes)}"

        node_ids = [node["node_id"] for node in article_nodes]
        assert len(set(node_ids)) == 2, f"Node IDs must be unique! Got: {node_ids}"

        # Verify their texts are separate and not merged together
        texts = [node["text"] for node in article_nodes]
        main_text_article = next(n for n in article_nodes if "Phạm vi điều chỉnh" in n["title"])
        appendix_article = next(n for n in article_nodes if "Công việc và địa điểm" in n["title"])

        assert "Công việc và địa điểm" not in main_text_article["text"]
        assert "Phạm vi điều chỉnh" not in appendix_article["text"]
