from __future__ import annotations

import asyncio
import io
import uuid
from datetime import UTC, datetime
from types import SimpleNamespace

import pytest
from docx import Document
from fastapi import HTTPException

from app.api import compare_contracts, draft_contract, review_contract
from app.schemas import (
    CompareContractRequest,
    DraftContractRequest,
    ReviewContractRequest,
    VerificationItem,
    VerificationReport,
)
from app.services.contract_analysis import (
    build_contract_diff,
    contract_retrieval_query,
    looks_like_contract,
)
from app.services.contract_documents import (
    MAX_DOCUMENT_BYTES,
    ContractDocumentError,
    extract_contract_document,
)
from app.services.contract_docx import (
    build_contract_docx,
    contract_download_filename,
    normalize_contract_plain_text,
)


def test_extract_contract_document_supports_utf8_and_docx_tables() -> None:
    plain = extract_contract_document(
        "HỢP ĐỒNG DỊCH VỤ\n\nĐiều 1. Phạm vi công việc".encode(),
        "hop-dong.txt",
    )
    assert plain.text.startswith("HỢP ĐỒNG DỊCH VỤ")
    assert plain.truncated is False

    document = Document()
    document.add_heading("HỢP ĐỒNG DỊCH VỤ", level=1)
    document.add_paragraph("BÊN A: CÔNG TY MINH AN")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Thời hạn"
    table.cell(0, 1).text = "12 tháng"
    document.add_paragraph("BÊN B: ÔNG NGUYỄN VĂN B")
    document.add_paragraph("Điều 1. Phạm vi công việc")
    output = io.BytesIO()
    document.save(output)

    extracted = extract_contract_document(output.getvalue(), "hop-dong.docx")
    assert "Điều 1. Phạm vi công việc" in extracted.text
    assert "Thời hạn\t12 tháng" in extracted.text
    assert extracted.text.index("BÊN A") < extracted.text.index("Thời hạn")
    assert extracted.text.index("Thời hạn") < extracted.text.index("BÊN B")
    assert extracted.text.index("BÊN B") < extracted.text.index("Điều 1")
    assert extracted.suggested_title == "HỢP ĐỒNG DỊCH VỤ"
    assert extracted.party_options == (
        "BÊN A — CÔNG TY MINH AN",
        "BÊN B — ÔNG NGUYỄN VĂN B",
    )


def test_extract_contract_document_rejects_unsupported_and_oversized_files() -> None:
    with pytest.raises(ContractDocumentError, match="Chỉ hỗ trợ"):
        extract_contract_document(b"not a contract", "contract.exe")
    with pytest.raises(ContractDocumentError, match="15 MB"):
        extract_contract_document(b"x" * (MAX_DOCUMENT_BYTES + 1), "contract.txt")


def test_contract_detection_and_retrieval_query_use_structure_not_full_body() -> None:
    short_request = "Soạn hợp đồng dịch vụ phát triển phần mềm."
    assert looks_like_contract(short_request) is False

    source = (
        "HỢP ĐỒNG DỊCH VỤ\n\n"
        "Điều 1. Phạm vi\nNội dung có thông tin riêng tư không được đưa vào truy vấn. " * 30
        + "\nĐiều 2. Thanh toán"
    )
    assert looks_like_contract(source) is True
    query = contract_retrieval_query("Rà soát hợp đồng", source)
    assert "Điều 1. Phạm vi" in query
    assert "thông tin riêng tư" not in query
    assert len(query) <= 1_600


def test_contract_diff_detects_added_deleted_and_modified_groups() -> None:
    original = "Điều 1. Phạm vi\n\nĐiều 2. Giá trị: 10 triệu\n\nĐiều 3. Bảo mật"
    revised = "Điều 1. Phạm vi\n\nĐiều 2. Giá trị: 12 triệu\n\nĐiều 4. Chấm dứt"
    result = build_contract_diff(original, revised)
    assert result["similarity"] < 100
    assert result["counts"]["modified"] >= 1
    assert result["changes"]


def test_contract_draft_is_plain_text_and_docx_uses_legal_page_layout() -> None:
    raw = (
        "```markdown\n# **HỢP ĐỒNG LAO ĐỘNG**\n\n"
        "## Điều 1. Công việc [S1]\n- Người lao động thực hiện [Công việc].\n\n"
        "NGƯỜI SỬ DỤNG LAO ĐỘNG | NGƯỜI LAO ĐỘNG\n"
        "(Ký, ghi rõ họ tên) | (Ký, ghi rõ họ tên)\n```"
    )
    plain = normalize_contract_plain_text(raw, "Hợp đồng lao động")
    assert plain.startswith("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    assert "Độc lập - Tự do - Hạnh phúc" in plain
    assert "Điều 1. Công việc" in plain
    assert "[Công việc]" in plain
    assert "[S1]" not in plain
    assert "**" not in plain
    assert "```" not in plain
    assert "|" not in plain

    data = build_contract_docx("Hợp đồng lao động", plain)
    document = Document(io.BytesIO(data))
    section = document.sections[0]
    assert section.page_width.cm == pytest.approx(21, abs=0.05)
    assert section.page_height.cm == pytest.approx(29.7, abs=0.05)
    assert section.left_margin.cm == pytest.approx(3, abs=0.05)
    assert section.right_margin.cm == pytest.approx(2, abs=0.05)
    assert any(
        "HỢP ĐỒNG LAO ĐỘNG" in paragraph.text for paragraph in document.paragraphs
    )
    assert any(
        "NGƯỜI SỬ DỤNG LAO ĐỘNG" in cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert contract_download_filename("Hợp đồng lao động") == "hop-dong-lao-dong.docx"


class _ArtifactDb:
    def __init__(self, user_id: uuid.UUID) -> None:
        self.user_id = user_id
        self.added: list[object] = []

    async def rollback(self) -> None:
        return None

    async def scalar(self, _: object) -> object:
        return self.user_id

    def add(self, value: object) -> None:
        self.added.append(value)

    async def commit(self) -> None:
        return None

    async def refresh(self, value: object) -> None:
        if getattr(value, "id", None) is None:
            value.id = uuid.uuid4()


class _FastContractRetrieval:
    def __init__(self) -> None:
        self.queries: list[str] = []

    async def retrieve(self, query: str) -> list[dict[str, object]]:
        self.queries.append(query)
        return [{
            "source_id": "S1",
            "doc_id": "law-1",
            "title": "Luật thử nghiệm",
            "citation": "100/2020/QH14",
            "text": "Nội dung pháp lý đã được kiểm chứng.",
        }]


class _Freshness:
    settings = SimpleNamespace(
        max_laws_verified_per_request=16,
        require_freshness_check=True,
    )

    async def verify_sources(
        self,
        _: list[dict[str, object]],
    ) -> tuple[VerificationReport, bool]:
        checked_at = datetime.now(UTC)
        return VerificationReport(
            checked=True,
            all_current=True,
            checked_at=checked_at,
            items=[
                VerificationItem(
                    code="100/2020/QH14",
                    title="Luật thử nghiệm",
                    status="IN_FORCE",
                    checked_at=checked_at,
                )
            ],
        ), False


def test_review_preserves_selected_party_and_accepts_contract_facts_without_statute_citation() -> None:
    user_id = uuid.uuid4()
    db = _ArtifactDb(user_id)
    retrieval = _FastContractRetrieval()

    class _Ai:
        async def complete_json(self, *_: object, **__: object) -> dict[str, object]:
            return {
                "summary": "Hợp đồng ghi thời hạn làm việc là 12 tháng.",
                "contract_type": "Hợp đồng lao động",
                "party_perspective": "Đánh giá cân bằng",
                "key_terms": [
                    {
                        "label": "Thời hạn",
                        "value": "12 tháng",
                        "assessment": "Được ghi trực tiếp trong hợp đồng.",
                    }
                ],
                "clause_reviews": [],
                "missing_clauses": [],
                "risks": [],
                "recommendations": ["Làm rõ ngày bắt đầu làm việc trong hợp đồng."],
            }

    result = asyncio.run(
        review_contract(
            ReviewContractRequest(
                title="  Hợp đồng lao động 2026  ",
                user_role="  NGƯỜI LAO ĐỘNG  ",
                text="HỢP ĐỒNG LAO ĐỘNG\r\n\r\nĐiều 1. Thời hạn làm việc là 12 tháng.",
            ),
            db,
            SimpleNamespace(id=user_id),
            SimpleNamespace(
                gemini_model="gemini-2.5-flash",
                message_encryption_key="",
                session_secret="test-session-secret-at-least-32-bytes",
            ),
            retrieval,
            _Freshness(),
            _Ai(),
        )
    )

    assert result["party_perspective"] == "NGƯỜI LAO ĐỘNG"
    assert result["contract_type"] == "Hợp đồng lao động"
    assert result["summary"].startswith("Hợp đồng ghi")
    assert db.added[0].title == "Hợp đồng lao động 2026"


def test_draft_accepts_full_contract_without_graph_route_or_mandatory_inline_citations() -> None:
    user_id = uuid.uuid4()
    db = _ArtifactDb(user_id)
    retrieval = _FastContractRetrieval()
    full_contract = (
        "HỢP ĐỒNG LAO ĐỘNG\n\nĐiều 1. Công việc\n"
        + ("Người lao động thực hiện công việc hiện có. " * 60)
        + "\n\nĐiều 2. Tiền lương"
    )

    class _Ai:
        async def complete(self, *_: object, **__: object) -> str:
            return "# **HỢP ĐỒNG LAO ĐỘNG**\n\nĐiều 1. Công việc đã được hoàn thiện [S1]."

    result = asyncio.run(
        draft_contract(
            DraftContractRequest(prompt=full_contract, template_name="Hợp đồng lao động"),
            db,
            SimpleNamespace(id=user_id),
            SimpleNamespace(
                gemini_model="gemini-2.5-flash",
                message_encryption_key="",
                session_secret="test-session-secret-at-least-32-bytes",
            ),
            retrieval,
            _Freshness(),
            _Ai(),
        )
    )
    assert result["draft"].startswith("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    assert "[S1]" not in result["draft"]
    assert "**" not in result["draft"]
    assert result["download_url"].endswith("/docx")
    assert len(retrieval.queries[0]) <= 1_600


def test_draft_rejects_non_labor_contract_template_before_retrieval() -> None:
    user_id = uuid.uuid4()
    db = _ArtifactDb(user_id)

    class _MustNotRun:
        def __getattr__(self, name: str) -> object:
            raise AssertionError(f"{name} must not run for an invalid template")

    with pytest.raises(HTTPException, match="liên quan trực tiếp đến lao động") as exc:
        asyncio.run(
            draft_contract(
                DraftContractRequest(
                    prompt="Soạn hợp đồng dịch vụ phát triển phần mềm.",
                    template_id="service",
                ),
                db,
                SimpleNamespace(id=user_id),
                SimpleNamespace(
                    gemini_model="gemini-2.5-flash",
                    message_encryption_key="",
                    session_secret="test-session-secret-at-least-32-bytes",
                ),
                _MustNotRun(),
                _MustNotRun(),
                _MustNotRun(),
            )
        )
    assert exc.value.status_code == 422


def test_draft_rejects_uploaded_contract_outside_labor_scope() -> None:
    user_id = uuid.uuid4()
    db = _ArtifactDb(user_id)
    service_contract = (
        "HỢP ĐỒNG DỊCH VỤ PHẦN MỀM\n\n"
        "Điều 1. Phạm vi dịch vụ\n"
        + ("Bên A thuê Bên B phát triển và bàn giao phần mềm. " * 30)
        + "\nĐiều 2. Phí dịch vụ"
    )

    class _MustNotRun:
        def __getattr__(self, name: str) -> object:
            raise AssertionError(f"{name} must not run for a non-labor contract")

    with pytest.raises(HTTPException, match="không có dấu hiệu") as exc:
        asyncio.run(
            draft_contract(
                DraftContractRequest(
                    prompt=service_contract,
                    template_id="employment",
                ),
                db,
                SimpleNamespace(id=user_id),
                SimpleNamespace(
                    gemini_model="gemini-2.5-flash",
                    message_encryption_key="",
                    session_secret="test-session-secret-at-least-32-bytes",
                ),
                _MustNotRun(),
                _MustNotRun(),
                _MustNotRun(),
            )
        )
    assert exc.value.status_code == 422


def test_identical_contracts_skip_retrieval_and_ai() -> None:
    user_id = uuid.uuid4()
    db = _ArtifactDb(user_id)

    class _MustNotRun:
        def __getattr__(self, name: str) -> object:
            raise AssertionError(f"{name} must not run for identical contracts")

    result = asyncio.run(
        compare_contracts(
            CompareContractRequest(
                original_text="Điều 1. Nội dung hợp đồng giống nhau.",
                revised_text="Điều 1. Nội dung hợp đồng giống nhau.",
            ),
            db,
            SimpleNamespace(id=user_id),
            SimpleNamespace(
                gemini_model="gemini-2.5-flash",
                message_encryption_key="",
                session_secret="test-session-secret-at-least-32-bytes",
            ),
            _MustNotRun(),
            _MustNotRun(),
            _MustNotRun(),
        )
    )
    assert result["similarity"] == 100
    assert result["differences"] == []
    assert result["sources"] == []
