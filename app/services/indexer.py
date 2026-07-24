from __future__ import annotations

import hashlib
import io
import re
import zipfile
from dataclasses import dataclass, replace
from typing import Any
from urllib.parse import urljoin, urlparse

import httpx
from bs4 import BeautifulSoup
from docx import Document
from fastapi.concurrency import run_in_threadpool
from neo4j import GraphDatabase
from pypdf import PdfReader
from sqlalchemy import delete, func, select, text as sql_text
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import Settings
from app.external_graphrag import (
    ExternalGraphRAGConfig,
    ensure_neo4j_schema,
    ensure_postgres_schema,
    postgres_connection,
    upsert_postgres_chunks,
)
from app.models import LegalChunk, LegalDocument, normalize_legal_document_code


ARTICLE_RE = re.compile(r"(?im)^\s*(Điều\s+\d+[a-zA-Z]?[\.:]?\s*[^\n]*)")
CLAUSE_RE = re.compile(r"(?m)^\s*(\d+)\.\s+")
HTTP_REDIRECT_STATUSES = {301, 302, 303, 307, 308}
MAX_LEGAL_DOCUMENT_BYTES = 25 * 1024 * 1024
MAX_LEGAL_DOCUMENT_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_LEGAL_DOCUMENT_TEXT_CHARS = 8_000_000
MAX_LEGAL_DOCUMENT_CHUNKS = 5_000
MAX_LEGAL_DOCUMENT_REDIRECTS = 5


class UnsafeLegalSourceError(ValueError):
    """The legal-source URL or response violates a non-bypassable safety rule."""


@dataclass(slots=True)
class LegalCandidate:
    code: str
    title: str
    url: str
    status: str
    issuer: str = ""
    external_doc_id: str | None = None
    replaces_code: str | None = None
    content: str | None = None


def _clean_text(value: str) -> str:
    value = value.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _bounded_join(parts: Any, *, max_chars: int) -> str:
    values: list[str] = []
    total = 0
    for raw_value in parts:
        value = str(raw_value or "")
        total += len(value)
        if total > max_chars:
            raise UnsafeLegalSourceError(
                "Văn bản sau giải nén vượt quá kích thước xử lý cho phép"
            )
        values.append(value)
    return "\n\n".join(values)


def _extract_pdf(content: bytes, max_chars: int) -> str:
    reader = PdfReader(io.BytesIO(content))
    return _bounded_join(
        (page.extract_text() or "" for page in reader.pages),
        max_chars=max_chars,
    )


def _extract_docx(content: bytes, max_chars: int) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            expanded_size = sum(item.file_size for item in archive.infolist())
    except (OSError, zipfile.BadZipFile) as exc:
        raise ValueError("Tệp DOCX từ nguồn chính thức không hợp lệ") from exc
    if expanded_size > MAX_LEGAL_DOCUMENT_UNCOMPRESSED_BYTES:
        raise UnsafeLegalSourceError(
            "Văn bản sau giải nén vượt quá kích thước xử lý cho phép"
        )
    document = Document(io.BytesIO(content))
    return _bounded_join(
        (paragraph.text for paragraph in document.paragraphs),
        max_chars=max_chars,
    )


def _extract_html(content: bytes, max_chars: int) -> str:
    soup = BeautifulSoup(content, "html.parser")
    for node in soup(["script", "style", "nav", "footer", "form", "noscript"]):
        node.decompose()
    root = soup.find("article") or soup.find("main") or soup.body or soup
    value = root.get_text("\n", strip=True)
    if len(value) > max_chars:
        raise UnsafeLegalSourceError(
            "Văn bản sau giải nén vượt quá kích thước xử lý cho phép"
        )
    return value


def normalize_law_code(value: str) -> str:
    return normalize_legal_document_code(value)


def _text_mentions_exact_law_code(value: str, law_code: str) -> bool:
    normalized_code = normalize_law_code(law_code)
    code_pattern = r"\s*".join(
        re.escape(character) for character in normalized_code
    )
    return bool(
        re.search(
            rf"(?<![0-9A-ZÀ-ỴĐ]){code_pattern}(?![0-9A-ZÀ-ỴĐ])",
            value.upper(),
        )
    )


def _require_official_https_url(
    value: str,
    allowed_domains: list[str],
    *,
    redirected: bool = False,
) -> None:
    try:
        parsed = urlparse(value)
        port = parsed.port
    except ValueError as exc:
        raise UnsafeLegalSourceError("URL nguồn pháp luật chính thức không hợp lệ") from exc
    if parsed.scheme != "https" or parsed.username or parsed.password:
        raise UnsafeLegalSourceError(
            "Từ chối tải văn bản từ URL không dùng HTTPS của nguồn chính thức"
        )
    if port is not None and port != 443:
        raise UnsafeLegalSourceError(
            "Từ chối tải văn bản từ cổng không chuẩn của nguồn chính thức"
        )
    host = (parsed.hostname or "").lower().removeprefix("www.")
    normalized_domains = {
        domain.strip().lower().removeprefix("www.")
        for domain in allowed_domains
        if domain.strip()
    }
    if not host or not any(
        host == domain or host.endswith(f".{domain}")
        for domain in normalized_domains
    ):
        if redirected:
            raise UnsafeLegalSourceError(
                "Từ chối lập chỉ mục vì URL sau chuyển hướng không thuộc nguồn chính thức"
            )
        raise UnsafeLegalSourceError(
            "Từ chối tải văn bản từ tên miền không thuộc danh sách nguồn chính thức"
        )


async def download_legal_text(
    url: str,
    timeout: int = 45,
    *,
    allowed_domains: list[str],
    max_bytes: int = MAX_LEGAL_DOCUMENT_BYTES,
    max_redirects: int = MAX_LEGAL_DOCUMENT_REDIRECTS,
) -> tuple[str, str]:
    headers = {"User-Agent": "VLegalAI/3.0 (+legal-document-refresh)"}
    current_url = url
    content = b""
    content_type = ""
    resolved_url = url
    async with httpx.AsyncClient(
        timeout=timeout,
        follow_redirects=False,
        headers=headers,
    ) as client:
        for redirect_count in range(max_redirects + 1):
            _require_official_https_url(
                current_url,
                allowed_domains,
                redirected=redirect_count > 0,
            )
            async with client.stream("GET", current_url) as response:
                if response.status_code in HTTP_REDIRECT_STATUSES:
                    location = response.headers.get("location")
                    if not location:
                        raise UnsafeLegalSourceError(
                            "Nguồn chính thức trả chuyển hướng không có Location"
                        )
                    if redirect_count >= max_redirects:
                        raise UnsafeLegalSourceError(
                            "Nguồn chính thức chuyển hướng quá số lần cho phép"
                        )
                    next_url = urljoin(str(response.url), location)
                    # Validate Location before issuing the next network request.
                    _require_official_https_url(
                        next_url,
                        allowed_domains,
                        redirected=True,
                    )
                    current_url = next_url
                    continue

                response.raise_for_status()
                resolved_url = str(response.url)
                _require_official_https_url(
                    resolved_url,
                    allowed_domains,
                    redirected=redirect_count > 0,
                )
                content_length = response.headers.get("content-length")
                if content_length:
                    try:
                        declared_length = int(content_length)
                    except ValueError:
                        declared_length = None
                    if declared_length is not None and declared_length > max_bytes:
                        raise UnsafeLegalSourceError(
                            "Văn bản vượt quá kích thước tải cho phép"
                        )
                content_type = response.headers.get("content-type", "").lower()
                buffer = bytearray()
                async for chunk in response.aiter_bytes():
                    buffer.extend(chunk)
                    if len(buffer) > max_bytes:
                        raise UnsafeLegalSourceError(
                            "Văn bản vượt quá kích thước tải cho phép"
                        )
                content = bytes(buffer)
                break
        else:  # pragma: no cover - loop exits via break or explicit redirect error
            raise UnsafeLegalSourceError(
                "Nguồn chính thức chuyển hướng quá số lần cho phép"
            )

    suffix = urlparse(resolved_url).path.lower()
    if "pdf" in content_type or suffix.endswith(".pdf"):
        text = await run_in_threadpool(
            _extract_pdf,
            content,
            MAX_LEGAL_DOCUMENT_TEXT_CHARS,
        )
    elif "word" in content_type or suffix.endswith(".docx"):
        text = await run_in_threadpool(
            _extract_docx,
            content,
            MAX_LEGAL_DOCUMENT_TEXT_CHARS,
        )
    else:
        text = await run_in_threadpool(
            _extract_html,
            content,
            MAX_LEGAL_DOCUMENT_TEXT_CHARS,
        )
    cleaned = _clean_text(text)
    if len(cleaned) < 500:
        raise ValueError("Nguồn chính thức không có đủ nội dung để tạo chỉ mục")
    return cleaned, resolved_url


def chunk_legal_text(candidate: LegalCandidate, text: str, version: int) -> list[dict[str, Any]]:
    if len(text) > MAX_LEGAL_DOCUMENT_TEXT_CHARS:
        raise UnsafeLegalSourceError(
            "Văn bản sau giải nén vượt quá kích thước xử lý cho phép"
        )
    matches: list[re.Match[str]] = []
    for match in ARTICLE_RE.finditer(text):
        matches.append(match)
        if len(matches) > MAX_LEGAL_DOCUMENT_CHUNKS:
            raise UnsafeLegalSourceError(
                "Văn bản tạo ra quá nhiều chunk để lập chỉ mục an toàn"
            )
    sections: list[tuple[str, str]] = []
    if matches:
        preamble = text[: matches[0].start()].strip()
        if preamble:
            sections.append((candidate.title, preamble))
        for index, match in enumerate(matches):
            end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
            sections.append((match.group(1).strip(), text[match.start() : end].strip()))
    else:
        size, overlap = 3500, 350
        cursor = 0
        while cursor < len(text):
            sections.append((f"Phần {len(sections) + 1}", text[cursor : cursor + size].strip()))
            cursor += size - overlap

    chunks: list[dict[str, Any]] = []
    ordinal = 0
    for heading, section in sections:
        parts = [section]
        if len(section) > 5500:
            clause_starts: list[re.Match[str]] = []
            for clause_start in CLAUSE_RE.finditer(section):
                clause_starts.append(clause_start)
                if len(clause_starts) > MAX_LEGAL_DOCUMENT_CHUNKS:
                    raise UnsafeLegalSourceError(
                        "Văn bản tạo ra quá nhiều chunk để lập chỉ mục an toàn"
                    )
            if len(clause_starts) > 1:
                parts = []
                for idx, match in enumerate(clause_starts):
                    end = clause_starts[idx + 1].start() if idx + 1 < len(clause_starts) else len(section)
                    parts.append(section[match.start() : end].strip())
        for part in parts:
            if not part:
                continue
            chunk_key = f"{candidate.code}:{version}:{ordinal}:{hashlib.sha256(part.encode('utf-8')).hexdigest()[:12]}"
            node_id = f"law:{candidate.code}:v{version}:section:{ordinal}"
            chunks.append(
                {
                    "external_chunk_id": chunk_key,
                    "node_id": node_id,
                    "chunk_type": "article" if heading.lower().startswith("điều") else "section",
                    "title": candidate.title,
                    "citation": f"{candidate.code} — {heading}",
                    "text": part,
                    "text_hash": hashlib.sha256(part.encode("utf-8")).hexdigest(),
                    "ordinal": ordinal,
                }
            )
            if len(chunks) > MAX_LEGAL_DOCUMENT_CHUNKS:
                raise UnsafeLegalSourceError(
                    "Văn bản tạo ra quá nhiều chunk để lập chỉ mục an toàn"
                )
            ordinal += 1
    return chunks


class LegalIndexer:
    def __init__(self, settings: Settings):
        self.settings = settings

    def _external_config(self) -> ExternalGraphRAGConfig:
        return ExternalGraphRAGConfig(
            neo4j_uri=self.settings.neo4j_uri,
            neo4j_user=self.settings.neo4j_user,
            neo4j_password=self.settings.neo4j_password,
            neo4j_database=self.settings.neo4j_database,
            database_url=self.settings.database_url,
            postgres_vector_size=self.settings.postgres_vector_size,
            embedding_model_path=self.settings.embedding_model_path,
            embedding_model_repo=self.settings.embedding_model_repo,
            embedding_model_revision=self.settings.embedding_model_revision,
            embedding_device=self.settings.embedding_device,
            embedding_batch_size=self.settings.embedding_batch_size,
            embedding_max_sequence_length=self.settings.embedding_max_sequence_length,
        )

    async def index_candidate(self, db: AsyncSession, candidate: LegalCandidate) -> LegalDocument:
        candidate = replace(
            candidate,
            code=normalize_law_code(candidate.code),
            replaces_code=(
                normalize_law_code(candidate.replaces_code)
                if candidate.replaces_code
                else None
            ),
        )
        _require_official_https_url(
            candidate.url,
            self.settings.official_legal_domains,
        )
        text, resolved_url = await download_legal_text(
            candidate.url,
            allowed_domains=self.settings.official_legal_domains,
        )
        _require_official_https_url(
            resolved_url,
            self.settings.official_legal_domains,
            redirected=resolved_url != candidate.url,
        )
        if not _text_mentions_exact_law_code(text, candidate.code):
            raise UnsafeLegalSourceError(
                "Nội dung tải từ nguồn chính thức không chứa đúng mã văn bản được yêu cầu"
            )

        # The database-level normalized unique index is the final guard. This
        # transaction-scoped lock also makes concurrent index refreshes for the
        # same normalized law code serialize instead of racing into that index.
        await db.execute(
            sql_text(
                "SELECT pg_advisory_xact_lock("
                "hashtextextended(:lock_key, 0)"
                ")"
            ),
            {"lock_key": f"legal-document:{candidate.code}"},
        )
        checksum = hashlib.sha256(text.encode("utf-8")).hexdigest()
        normalized_db_code = func.upper(
            func.regexp_replace(
                func.btrim(LegalDocument.code),
                "[[:space:]]+",
                "",
                "g",
            )
        )
        document = await db.scalar(
            select(LegalDocument).where(normalized_db_code == candidate.code)
        )
        if document and document.checksum == checksum:
            document.status = candidate.status
            document.source_url = resolved_url
            existing_chunks = list(
                (
                    await db.scalars(
                        select(LegalChunk)
                        .where(
                            LegalChunk.document_id == document.id,
                            LegalChunk.version == document.version,
                        )
                        .order_by(LegalChunk.ordinal)
                    )
                ).all()
            )
            if not existing_chunks:
                for chunk in chunk_legal_text(candidate, text, document.version):
                    row = LegalChunk(
                        document_id=document.id,
                        version=document.version,
                        **chunk,
                    )
                    db.add(row)
                    existing_chunks.append(row)
                await db.flush()
            await run_in_threadpool(
                self._sync_external,
                document,
                existing_chunks,
                candidate.replaces_code,
            )
            return document
        if not document:
            document = LegalDocument(
                code=candidate.code,
                title=candidate.title,
                issuer=candidate.issuer or None,
                external_doc_id=candidate.external_doc_id,
                source_url=resolved_url,
                official_domain=urlparse(resolved_url).netloc.lower(),
                status=candidate.status,
                checksum=checksum,
                version=1,
            )
            db.add(document)
            await db.flush()
        else:
            if document.checksum:
                document.version += 1
            document.title = candidate.title or document.title
            document.external_doc_id = document.external_doc_id or candidate.external_doc_id
            document.source_url = resolved_url
            document.official_domain = urlparse(resolved_url).netloc.lower()
            document.status = candidate.status
            document.checksum = checksum

        chunks = chunk_legal_text(candidate, text, document.version)
        await db.execute(
            delete(LegalChunk).where(LegalChunk.document_id == document.id, LegalChunk.version == document.version)
        )
        rows: list[LegalChunk] = []
        for chunk in chunks:
            row = LegalChunk(
                document_id=document.id,
                version=document.version,
                **chunk,
            )
            db.add(row)
            rows.append(row)
        await db.flush()
        await run_in_threadpool(self._sync_external, document, rows, candidate.replaces_code)
        return document

    @staticmethod
    def _link_replacement(
        session: Any,
        document: LegalDocument,
        replaces_code: str,
    ) -> None:
        session.run(
            """
            MATCH (new:LegalNode)
            WHERE new.node_id = $new_id
               OR toUpper(coalesce(new.code, '')) = toUpper($new_code)
               OR toUpper(coalesce(new.number, '')) = toUpper($new_code)
            WITH new,
                 CASE WHEN new.node_id = $new_id THEN 0 ELSE 1 END AS priority
            ORDER BY priority
            LIMIT 1
            SET new.status = $new_status
            WITH new
            MATCH (old:LegalNode)
            WHERE old.node_id <> new.node_id
              AND old.node_type IN ['document', 'VănBản']
              AND (
                    toUpper(coalesce(old.code, '')) = toUpper($old_code)
                 OR toUpper(coalesce(old.number, '')) = toUpper($old_code)
              )
            SET old.status = 'REPLACED'
            MERGE (new)-[:REPLACES]->(old)
            WITH collect(DISTINCT old.doc_id) AS old_doc_ids
            MATCH (old_chunk:LegalChunk)
            WHERE old_chunk.doc_id IN old_doc_ids
               OR old_chunk.citation CONTAINS $old_parenthesized_code
               OR old_chunk.citation STARTS WITH $old_code_prefix
            SET old_chunk.law_status = 'REPLACED'
            """,
            new_id=f"law:{document.code}:v{document.version}",
            new_code=document.code,
            new_status=document.status,
            old_code=replaces_code,
            old_parenthesized_code=f"({replaces_code})",
            old_code_prefix=f"{replaces_code} —",
        )

    def _sync_replacement_state(
        self,
        document: LegalDocument,
        replaces_code: str,
    ) -> None:
        config = self._external_config()
        if not config.neo4j_password:
            return
        driver = GraphDatabase.driver(
            config.neo4j_uri,
            auth=(config.neo4j_user, config.neo4j_password),
        )
        try:
            ensure_neo4j_schema(driver, config.neo4j_database)
            with driver.session(database=config.neo4j_database) as session:
                with session.begin_transaction() as transaction:
                    self._link_replacement(
                        transaction,
                        document,
                        replaces_code,
                    )
        finally:
            driver.close()

    @staticmethod
    def _mark_postgres_replaced(
        config: ExternalGraphRAGConfig,
        replaces_code: str,
    ) -> None:
        with postgres_connection(config) as connection:
            with connection.transaction():
                with connection.cursor() as cursor:
                    cursor.execute(
                        """
                        UPDATE graphrag_chunk
                        SET law_status = 'REPLACED',
                            updated_at = now()
                        WHERE upper(
                            regexp_replace(
                                btrim(law_code),
                                '[[:space:]]+',
                                '',
                                'g'
                            )
                        ) = %(law_code)s
                        """,
                        {"law_code": normalize_law_code(replaces_code)},
                    )

    def _sync_external(
        self, document: LegalDocument, chunks: list[LegalChunk], replaces_code: str | None
    ) -> None:
        config = self._external_config()
        if config.postgres_ready:
            ensure_postgres_schema(config)
            upsert_postgres_chunks(
                [
                    {
                        "chunk_id": chunk.external_chunk_id,
                        "doc_id": document.external_doc_id or str(document.id),
                        "node_id": chunk.node_id,
                        "chunk_type": chunk.chunk_type,
                        "title": chunk.title,
                        "path_label": chunk.citation,
                        "citation": chunk.citation,
                        "text": chunk.text,
                        "token_count": len(chunk.text.split()),
                        "ordinal": chunk.ordinal,
                        "source_url": document.source_url,
                        "law_code": document.code,
                        "law_status": document.status,
                        "law_version": document.version,
                    }
                    for chunk in chunks
                ],
                config,
            )

        if config.neo4j_password:
            driver = GraphDatabase.driver(config.neo4j_uri, auth=(config.neo4j_user, config.neo4j_password))
            try:
                ensure_neo4j_schema(driver, config.neo4j_database)
                with driver.session(database=config.neo4j_database) as session:
                    with session.begin_transaction() as transaction:
                        transaction.run(
                            """
                            MERGE (d:LegalNode {node_id: $node_id})
                            SET d.node_type='document', d.doc_id=$doc_id, d.code=$code,
                                d.title=$title, d.status=$status, d.source_url=$source_url, d.version=$version
                            """,
                            node_id=f"law:{document.code}:v{document.version}",
                            doc_id=document.external_doc_id or str(document.id),
                            code=document.code,
                            title=document.title,
                            status=document.status,
                            source_url=document.source_url,
                            version=document.version,
                        )
                        transaction.run(
                            """
                            UNWIND $rows AS row
                            MERGE (n:LegalNode {node_id: row.node_id})
                            SET n.node_type=row.chunk_type, n.doc_id=$doc_id, n.title=row.citation,
                                n.text=row.text, n.ordinal=row.ordinal
                            MERGE (c:LegalChunk {chunk_id: row.chunk_id})
                            SET c.node_id=row.node_id, c.doc_id=$doc_id, c.chunk_type=row.chunk_type,
                                c.title=$title, c.citation=row.citation, c.text=row.text,
                                c.ordinal=row.ordinal, c.source_url=$source_url, c.version=$version,
                                c.law_code=$code, c.law_status=$status, c.law_version=$version
                            MERGE (c)-[:CHUNK_OF]->(n)
                            MERGE (n)-[:BELONGS_TO]->(d)
                            """,
                            rows=[
                                {
                                    "node_id": chunk.node_id,
                                    "chunk_id": chunk.external_chunk_id,
                                    "chunk_type": chunk.chunk_type,
                                    "citation": chunk.citation,
                                    "text": chunk.text,
                                    "ordinal": chunk.ordinal,
                                }
                                for chunk in chunks
                            ],
                            doc_id=document.external_doc_id or str(document.id),
                            title=document.title,
                            source_url=document.source_url,
                            code=document.code,
                            version=document.version,
                            status=document.status,
                        )
            finally:
                driver.close()

        # Publish replacement state only after every enabled backend has the
        # new document and chunks. There is no distributed transaction across
        # the app database, PostgreSQL index connection, and Neo4j; keeping the
        # old law current on a partial failure is the fail-closed behavior.
        if replaces_code:
            if config.postgres_ready:
                self._mark_postgres_replaced(config, replaces_code)
            self._sync_replacement_state(document, replaces_code)
