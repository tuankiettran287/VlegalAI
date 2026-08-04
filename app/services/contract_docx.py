from __future__ import annotations

import io
import re
import unicodedata

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

_CITATION_RE = re.compile(r"\s*\[S\d+\]", re.IGNORECASE)
_MARKDOWN_LINK_RE = re.compile(r"\[([^\]]+)]\((?:[^()]|\([^)]*\))+\)")
_MARKDOWN_TABLE_SEPARATOR_RE = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$"
)
_ARTICLE_HEADING_RE = re.compile(r"^Điều\s+\d+[a-zA-ZđĐ]?(?:[.:]|\s|$)", re.IGNORECASE)
_PREFACE_RE = re.compile(
    r"^(?:dưới đây|sau đây)\s+là\s+(?:bản\s+)?(?:dự thảo|hợp đồng)|"
    r"^tôi sẽ\s+(?:soạn|hoàn thiện)",
    re.IGNORECASE,
)

NATIONAL_HEADING = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
NATIONAL_MOTTO = "Độc lập - Tự do - Hạnh phúc"


def _fold(value: str) -> str:
    normalized = unicodedata.normalize(
        "NFKD", value.replace("Đ", "D").replace("đ", "d")
    )
    without_marks = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    return re.sub(r"\s+", " ", without_marks).strip().casefold()


def _clean_line(line: str) -> str:
    value = line.replace("\u00a0", " ").strip()
    value = re.sub(r"^\s{0,3}#{1,6}\s+", "", value)
    value = _MARKDOWN_LINK_RE.sub(r"\1", value)
    value = _CITATION_RE.sub("", value)
    value = value.replace("**", "").replace("__", "").replace("`", "")
    value = re.sub(r"^\s*[-*+]\s+", "• ", value)
    if value.count("|") == 1:
        value = re.sub(r"\s*\|\s*", "\t", value)
    return re.sub(r" +", " ", value).strip()


def normalize_contract_plain_text(content: str, title: str) -> str:
    """Convert an AI draft into safe, plain legal-document text.

    The returned value deliberately contains no Markdown. It is suitable for
    both the UI document preview and deterministic DOCX rendering.
    """

    lines: list[str] = []
    for raw_line in (
        content.replace("\x00", "")
        .replace("\r\n", "\n")
        .replace("\r", "\n")
        .split("\n")
    ):
        if raw_line.strip().startswith("```"):
            continue
        if _MARKDOWN_TABLE_SEPARATOR_RE.match(raw_line):
            continue
        line = _clean_line(raw_line)
        if not line:
            if lines and lines[-1] != "":
                lines.append("")
            continue
        if not lines and _PREFACE_RE.search(line):
            continue
        lines.append(line)

    while lines and not lines[-1]:
        lines.pop()

    title_text = (title or "Hợp đồng lao động").strip()
    body_lines: list[str] = []
    document_title = ""
    for line in lines:
        folded = _fold(line)
        if folded in {_fold(NATIONAL_HEADING), _fold(NATIONAL_MOTTO)}:
            continue
        if (
            not document_title
            and len(line) <= 180
            and _is_upper_section(line)
            and any(
                signal in folded for signal in ("hop dong", "thoa thuan", "phu luc")
            )
        ):
            document_title = line
            continue
        body_lines.append(line)

    normalized_lines = [
        NATIONAL_HEADING,
        NATIONAL_MOTTO,
        "",
        (document_title or title_text).upper(),
        "",
        *body_lines,
    ]
    compacted: list[str] = []
    for line in normalized_lines:
        if line or (compacted and compacted[-1]):
            compacted.append(line)
    return "\n".join(compacted).strip()


def _set_cell_borderless(cell: object) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_run_font(
    run: object,
    *,
    size: float = 13,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_page_number(paragraph: object) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    _set_run_font(run, size=10)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def _is_upper_section(line: str) -> bool:
    letters = [ch for ch in line if ch.isalpha()]
    return (
        bool(letters) and len(line) <= 120 and all(not ch.islower() for ch in letters)
    )


def build_contract_docx(title: str, content: str) -> bytes:
    """Render a labor-contract draft as an A4 Vietnamese legal document."""

    normalized = normalize_contract_plain_text(content, title)
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(3)

    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    title_folded = _fold(title)
    for line in normalized.splitlines():
        if not line:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            continue

        signature_cells = [part.strip() for part in re.split(r"\t+|\s+\|\s+", line)]
        if len(signature_cells) == 2 and any(
            signal in _fold(line)
            for signal in (
                "nguoi su dung lao dong",
                "nguoi lao dong",
                "ky, ghi ro ho ten",
            )
        ):
            table = document.add_table(rows=1, cols=2)
            table.autofit = True
            for cell, value in zip(table.rows[0].cells, signature_cells, strict=True):
                _set_cell_borderless(cell)
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(value)
                _set_run_font(
                    run,
                    bold="nguoi" in _fold(value) and "ky" not in _fold(value),
                    italic="ky" in _fold(value),
                )
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(3)
        folded = _fold(line)

        if folded == _fold(NATIONAL_HEADING):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line)
            _set_run_font(run, bold=True)
        elif folded == _fold(NATIONAL_MOTTO):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line)
            _set_run_font(run, bold=True)
            run.underline = True
        elif (
            folded == title_folded
            or "hop dong lao dong" in folded
            or folded.startswith("phu luc hop dong lao dong")
        ):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(line.upper())
            _set_run_font(run, size=16, bold=True)
        elif _ARTICLE_HEADING_RE.match(line):
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(line)
            _set_run_font(run, bold=True)
        elif line.startswith("• "):
            paragraph.style = document.styles["List Bullet"]
            paragraph.paragraph_format.left_indent = Cm(0.65)
            run = paragraph.add_run(line[2:].strip())
            _set_run_font(run)
        elif folded.startswith("can cu"):
            run = paragraph.add_run(line)
            _set_run_font(run, italic=True)
        elif _is_upper_section(line):
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(line)
            _set_run_font(run, bold=True)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(line)
            _set_run_font(run)

    _add_page_number(section.footer.paragraphs[0])
    properties = document.core_properties
    properties.title = title
    properties.subject = "Hợp đồng lao động do VLegal AI hỗ trợ soạn thảo"
    properties.author = "VLegal AI"

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def contract_download_filename(title: str) -> str:
    folded = unicodedata.normalize("NFKD", title.replace("Đ", "D").replace("đ", "d"))
    ascii_title = "".join(ch for ch in folded if not unicodedata.combining(ch))
    safe = re.sub(r"[^A-Za-z0-9]+", "-", ascii_title).strip("-").lower()
    return f"{safe or 'hop-dong-lao-dong'}.docx"
