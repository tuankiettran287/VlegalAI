from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

MAX_DOCUMENT_BYTES = 15 * 1024 * 1024
MAX_EXTRACTED_CHARS = 120_000
MAX_PDF_PAGES = 250
MAX_DOCX_ENTRIES = 5_000
MAX_DOCX_UNCOMPRESSED_BYTES = 40 * 1024 * 1024
SUPPORTED_DOCUMENT_EXTENSIONS = {".docx", ".md", ".pdf", ".txt"}


class ContractDocumentError(ValueError):
    """Raised when an uploaded contract cannot be read safely."""


@dataclass(frozen=True, slots=True)
class ExtractedContractDocument:
    filename: str
    text: str
    original_chars: int
    truncated: bool
    page_count: int | None = None
    suggested_title: str | None = None
    party_options: tuple[str, ...] = ()


_CONTRACT_TITLE_RE = re.compile(
    r"\b(?:hợp\s+đồng|thỏa\s+thuận|thoả\s+thuận|phụ\s+lục)\b",
    re.IGNORECASE,
)
_PARTY_RE = re.compile(
    r"^(?P<label>"
    r"bên\s+(?:[a-z0-9]+|sử\s+dụng\s+lao\s+động|thuê|cho\s+thuê|mua|bán|giao|nhận|cung\s+cấp\s+dịch\s+vụ)"
    r"|người\s+sử\s+dụng\s+lao\s+động"
    r"|người\s+lao\s+động"
    r"|bên\s+đặt\s+hàng"
    r"|bên\s+nhận\s+việc"
    r")\s*(?:[:\-–—]|\(|$)\s*(?P<detail>.*)$",
    re.IGNORECASE,
)


def _single_line(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip(" \t:|–—-")


def suggest_contract_title(text: str, filename: str = "") -> str | None:
    """Return a human-readable title without calling an LLM."""

    candidates: list[tuple[int, int, str]] = []
    for index, raw_line in enumerate(text.splitlines()[:160]):
        line = _single_line(raw_line)
        if not 5 <= len(line) <= 180 or not _CONTRACT_TITLE_RE.search(line):
            continue
        word_count = len(line.split())
        if word_count > 22:
            continue
        upper_letters = [character for character in line if character.isalpha()]
        uppercase = bool(upper_letters) and all(
            not character.islower() for character in upper_letters
        )
        starts_as_title = bool(
            re.match(
                r"^(?:hợp\s+đồng|thỏa\s+thuận|thoả\s+thuận|phụ\s+lục)",
                line,
                re.IGNORECASE,
            )
        )
        score = (5 if starts_as_title else 0) + (3 if uppercase else 0) - min(index, 40)
        candidates.append((score, -index, line))
    if candidates:
        return max(candidates)[2]

    stem = _single_line(re.sub(r"[_-]+", " ", Path(filename).stem))
    if stem and stem.casefold() not in {"document", "hop dong", "hợp đồng"}:
        return stem[:160]
    return None


def extract_contract_parties(text: str) -> tuple[str, ...]:
    """Extract party labels/names so the user can choose the review perspective."""

    parties: list[str] = []
    seen: set[str] = set()
    for raw_line in text.splitlines()[:500]:
        # A DOCX table row may contain both parties. Inspect each cell as well
        # as the full row, without losing the human-readable row in `text`.
        segments = re.split(r"\t+|\s+\|\s+", raw_line)
        for raw_segment in segments:
            segment = _single_line(raw_segment)
            if not segment or len(segment) > 220:
                continue
            match = _PARTY_RE.match(segment)
            if not match:
                continue
            label = _single_line(match.group("label"))
            detail = _single_line(match.group("detail"))
            if detail.casefold().startswith(("sau đây", "đại diện", "địa chỉ", "mã số")):
                detail = ""
            display = label if not detail else f"{label} — {detail[:120]}"
            identity = re.sub(r"[^0-9a-zđ]+", " ", label.casefold()).strip()
            if identity in seen:
                continue
            seen.add(identity)
            parties.append(display)
            if len(parties) >= 8:
                return tuple(parties)
    return tuple(parties)


def _normalize_text(value: str) -> str:
    normalized = (
        value.replace("\x00", "")
        .replace("\x0c", "\n")
        .replace("\r\n", "\n")
        .replace("\r", "\n")
    )
    normalized = "\n".join(line.rstrip() for line in normalized.splitlines())
    return re.sub(r"\n{4,}", "\n\n\n", normalized).strip()


def _bounded_text(value: str, filename: str, page_count: int | None = None) -> ExtractedContractDocument:
    normalized = _normalize_text(value)
    if len(normalized) < 20:
        raise ContractDocumentError(
            "Không đọc được nội dung văn bản. Nếu đây là bản scan, hãy dùng file PDF có lớp chữ hoặc DOCX."
        )
    original_chars = len(normalized)
    return ExtractedContractDocument(
        filename=filename,
        text=normalized[:MAX_EXTRACTED_CHARS],
        original_chars=original_chars,
        truncated=original_chars > MAX_EXTRACTED_CHARS,
        page_count=page_count,
        suggested_title=suggest_contract_title(normalized, filename),
        party_options=extract_contract_parties(normalized),
    )


def _decode_plain_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1258"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ContractDocumentError("File văn bản phải dùng bảng mã UTF-8, UTF-16 hoặc tiếng Việt Windows.")


def _extract_docx(data: bytes) -> tuple[str, None]:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            entries = archive.infolist()
            if (
                len(entries) > MAX_DOCX_ENTRIES
                or sum(entry.file_size for entry in entries) > MAX_DOCX_UNCOMPRESSED_BYTES
            ):
                raise ContractDocumentError("File DOCX vượt quá giới hạn giải nén an toàn.")
    except zipfile.BadZipFile as exc:
        raise ContractDocumentError("File DOCX không hợp lệ hoặc đã bị hỏng.") from exc

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ContractDocumentError("Không thể đọc file DOCX này.") from exc

    # `document.paragraphs` and `document.tables` are separate collections.
    # Concatenating them moves every table to the end of the contract and
    # destroys clause/party context. Iterate the underlying body instead.
    blocks: list[str] = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, document)
            if paragraph.text.strip():
                blocks.append(paragraph.text)
            continue
        if not isinstance(child, CT_Tbl):
            continue
        table = Table(child, document)
        for row in table.rows:
            cells = [_single_line(cell.text) for cell in row.cells]
            if any(cells):
                blocks.append("\t".join(cells))
    return "\n\n".join(blocks), None


def _extract_pdf(data: bytes) -> tuple[str, int]:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ContractDocumentError("File PDF không hợp lệ hoặc đã bị hỏng.") from exc
    if reader.is_encrypted:
        raise ContractDocumentError("Không hỗ trợ PDF đang được đặt mật khẩu.")
    page_count = len(reader.pages)
    if page_count > MAX_PDF_PAGES:
        raise ContractDocumentError(f"PDF chỉ được tối đa {MAX_PDF_PAGES} trang.")

    pages: list[str] = []
    try:
        for page in reader.pages:
            try:
                text = page.extract_text(extraction_mode="layout") or ""
            except TypeError:
                # Backwards compatibility for older pypdf releases.
                text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
            if sum(len(item) for item in pages) > MAX_EXTRACTED_CHARS:
                break
    except Exception as exc:
        raise ContractDocumentError("Không thể trích xuất chữ từ PDF này.") from exc
    return "\n\n".join(pages), page_count


def extract_contract_document(
    data: bytes,
    filename: str,
    content_type: str | None = None,
) -> ExtractedContractDocument:
    if not data:
        raise ContractDocumentError("File tải lên đang trống.")
    if len(data) > MAX_DOCUMENT_BYTES:
        raise ContractDocumentError("File hợp đồng chỉ được tối đa 15 MB.")

    safe_filename = Path(filename or "hop-dong").name[:255]
    extension = Path(safe_filename).suffix.casefold()
    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise ContractDocumentError("Chỉ hỗ trợ file PDF, DOCX, TXT hoặc Markdown.")

    if extension in {".txt", ".md"}:
        text, page_count = _decode_plain_text(data), None
    elif extension == ".docx":
        text, page_count = _extract_docx(data)
    else:
        text, page_count = _extract_pdf(data)
    return _bounded_text(text, safe_filename, page_count)
