from __future__ import annotations

import hashlib
import json
import logging
import math
import os
import re
import sqlite3
import unicodedata
from array import array
from collections import OrderedDict
from pathlib import Path
from typing import Any, Iterable

from docx import Document

from app import legal_ontology as onto
from app.services.embedding_checkpoint import (
    EmbeddingCheckpointRecord,
    PostgresEmbeddingCheckpoint,
    embedding_content_hash,
)
from app.services.embeddings import EmbeddingConfig, get_embedding_service


logger = logging.getLogger(__name__)
PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DATA_DIR = PROJECT_ROOT / "Data (1)"
DEFAULT_STORAGE_DIR = PROJECT_ROOT / "storage" / "graphrag"
DEFAULT_DB_PATH = DEFAULT_STORAGE_DIR / "legal_graphrag.sqlite"

SYSTEM_DOC_ID = "he-thong"

CHUNK_WINDOW_WORDS = 360
CHUNK_OVERLAP_WORDS = 70

#: A single term may not generate more MENTIONS edges than this, otherwise a
#: ubiquitous phrase such as "người lao động" would connect to half the graph.
MAX_MENTION_EDGES_PER_TERM = 220

VN_WORD_RE = re.compile(r"[0-9A-Za-zÀ-ỹĐđ]+", re.UNICODE)
CHAPTER_RE = re.compile(r"^Chương\s+([IVXLCDM]+|\d+)(?:[\.\s:-]+(.+))?$", re.IGNORECASE)
SECTION_RE = re.compile(r"^Mục\s+([IVXLCDM]+|\d+)(?:[\.\s:-]+(.+))?$", re.IGNORECASE)
ARTICLE_RE = re.compile(r"^Điều\s+(\d+[a-zA-Z]?)\s*[\.:]\s*(.+)$", re.IGNORECASE)
CLAUSE_RE = re.compile(r"^(\d{1,3})\.\s+(.+)$")
POINT_RE = re.compile(r"^([a-zđ](?:\d+)?)\)\s+(.+)$", re.IGNORECASE)
ARTICLE_REF_RE = re.compile(
    r"(?:(?:điểm)\s+([a-zđ](?:\d+)?)\s+)?"
    r"(?:(?:khoản)\s+(\d{1,3})\s+)?"
    r"Điều\s+(\d+[a-zA-Z]?)",
    re.IGNORECASE,
)

#: "1. Người lao động là người làm việc..." inside a "Giải thích từ ngữ" article.
DEFINITION_RE = re.compile(
    r"^\s*(?:\d{1,3}\.\s*)?(?P<term>[^.;:]{3,80}?)\s+(?:là|được hiểu là|được gọi là)\s+(?P<body>.{20,}?)\s*$",
    re.IGNORECASE | re.DOTALL,
)
DEFINITION_ARTICLE_RE = re.compile(r"giải thích (?:từ ngữ|thuật ngữ)", re.IGNORECASE)

#: "Phạt tiền từ 5.000.000 đồng đến 10.000.000 đồng"
FINE_RANGE_RE = re.compile(
    r"[Pp]hạt\s+tiền\s+từ\s+([\d\.]+)\s*đồng\s+đến\s+([\d\.]+)\s*đồng",
    re.IGNORECASE,
)
#: "Phạt tiền 10.000.000 đồng" / "phạt tiền từ 2.000.000 đồng trở lên"
FINE_SINGLE_RE = re.compile(r"[Pp]hạt\s+tiền\s+(?:từ\s+)?([\d\.]{5,})\s*đồng", re.IGNORECASE)
MONEY_RE = re.compile(r"\b(\d{1,3}(?:\.\d{3})+)\s*đồng\b")
PERCENT_RE = re.compile(r"\b(\d{1,3}(?:[,.]\d+)?)\s*%")
DURATION_RE = re.compile(
    r"\b(\d{1,4})\s*(giờ|ngày làm việc|ngày|tuần|tháng|năm)\b",
    re.IGNORECASE,
)
EFFECTIVE_DATE_RE = re.compile(
    r"có hiệu lực (?:thi hành|thực hiện)?\s*(?:kể\s+)?từ ngày\s+(\d{1,2})\s*(?:tháng|/|-)\s*(\d{1,2})\s*(?:năm|/|-)\s*(\d{4})",
    re.IGNORECASE,
)
REGION_WAGE_RE = re.compile(
    r"Vùng\s+([IVX]+)\s*\|\s*([\d\.]+)\s*\|\s*([\d\.]+)\b",
    re.IGNORECASE,
)

DOC_TYPE_PREFIXES: tuple[tuple[str, str], ...] = (
    ("bo-luat", "Bộ luật"),
    ("nghi-dinh", "Nghị định"),
    ("nghi-quyet", "Nghị quyết"),
    ("thong-tu", "Thông tư"),
    ("quyet-dinh", "Quyết định"),
    ("luat", "Luật"),
)


def normalize_space(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\t", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_block(text: str) -> str:
    """Normalise whitespace but keep line breaks.

    Table rows only mean something as rows: collapsing them onto one line turned
    "Vùng I | 5.310.000 | 25.500" into an unparseable run-on, so wage brackets
    and fee scales could never be read back out of the index.
    """

    lines = [normalize_space(line) for line in (text or "").splitlines()]
    return "\n".join(line for line in lines if line)


def strip_accents(text: str) -> str:
    text = text.replace("Đ", "D").replace("đ", "d")
    return "".join(
        c for c in unicodedata.normalize("NFD", text) if unicodedata.category(c) != "Mn"
    )


def slugify(text: str, fallback: str = "item") -> str:
    text = text.replace("Đ", "DD").replace("đ", "dd")
    text = strip_accents(text).lower()
    text = re.sub(r"[^a-z0-9]+", "-", text).strip("-")
    return text or fallback


def is_separator(text: str) -> bool:
    return bool(re.fullmatch(r"[_=\-\s\.]{3,}", text or ""))


def uppercase_ratio(text: str) -> float:
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return 0.0
    return sum(1 for c in letters if c.isupper()) / len(letters)


def is_heading_title(text: str) -> bool:
    if not text or len(text) > 220:
        return False
    if ARTICLE_RE.match(text) or CHAPTER_RE.match(text) or SECTION_RE.match(text):
        return False
    if CLAUSE_RE.match(text) or POINT_RE.match(text):
        return False
    return uppercase_ratio(text) >= 0.55 or text[:1].isupper()


def token_count(text: str) -> int:
    return len(VN_WORD_RE.findall(text))


def embedding_text_windows(text: str) -> list[str]:
    """Bound structural chunks so Vertex never silently drops their tail."""

    if token_count(text) <= CHUNK_WINDOW_WORDS + 80:
        return [text]

    raw_words = text.split()
    step = max(80, CHUNK_WINDOW_WORDS - CHUNK_OVERLAP_WORDS)
    windows: list[str] = []
    for start in range(0, len(raw_words), step):
        window = raw_words[start : start + CHUNK_WINDOW_WORDS]
        if len(window) < 80:
            break
        windows.append(" ".join(window))
    return windows or [text]


#: Letterhead lines that precede the real title of a Vietnamese legal document.
MASTHEAD_RE = re.compile(
    r"^(quoc hoi|chinh phu|uy ban thuong vu quoc hoi|bo (?!luat\b)[a-z]|cong hoa xa hoi|doc lap|"
    r"so:|luat so:|bo luat so:|nghi dinh so:|thong tu so:|nghi quyet so:|ha noi, ngay|"
    r"[a-z\s]*, ngay \d)",
    re.IGNORECASE,
)


def smart_doc_title(lines: list[str], filename: str) -> str:
    selected: list[str] = []
    for line in lines[:16]:
        if is_separator(line):
            continue
        # Table rows and mastheads carry no title information.
        if "|" in line or MASTHEAD_RE.match(strip_accents(line).lower()):
            continue
        if re.match(r"^(Căn cứ|Theo đề nghị|Quốc hội ban hành|Chính phủ ban hành)", line, re.I):
            break
        if CHAPTER_RE.match(line) or ARTICLE_RE.match(line):
            break
        selected.append(line)
        if len(selected) >= 3:
            break
    title = normalize_space(" ".join(selected)) or Path(filename).stem.replace("-", " ")
    if uppercase_ratio(title) > 0.85:
        title = title.title()
    return title


def detect_code(filename: str, lines: list[str]) -> str:
    stem = Path(filename).stem
    normalized = stem.replace("_", "-")
    m = re.search(r"(\d+)-(\d{4})-([A-ZĐ0-9]+(?:-[A-ZĐ0-9]+)*)", normalized, re.I)
    if m:
        return f"{m.group(1)}/{m.group(2)}/{m.group(3).upper()}"
    m = re.match(r"(\d+)[-_](VBHN-[A-ZĐ]+)", stem, re.I)
    if m:
        return f"{m.group(1)}/{m.group(2).upper()}"
    joined = " ".join(lines[:20])
    m = re.search(r"(\d+)\s*/\s*(\d{4})\s*/\s*([A-ZĐ0-9]+(?:-[A-ZĐ0-9]+)*)", joined, re.I)
    if m:
        return f"{m.group(1)}/{m.group(2)}/{m.group(3).upper()}"
    return stem


def detect_doc_type(filename: str, title: str, code: str = "") -> str:
    """Classify a document.

    The filename prefix is authoritative: the body of a decree routinely quotes
    "Bộ luật Lao động" in its opening lines, so scoring on body text alone
    misfiled most Nghị định / Thông tư as "Bộ luật" and broke the HƯỚNG_DẪN
    edges that depend on the type pair.
    """

    slug = slugify(Path(filename).stem)
    for prefix, doc_type in DOC_TYPE_PREFIXES:
        if slug.startswith(prefix):
            return doc_type

    code_upper = strip_accents(code or "").upper()
    if "VBHN" in code_upper:
        return "Văn bản hợp nhất"
    if "ND-CP" in code_upper:
        return "Nghị định"
    if code_upper.startswith("TT-") or "-TT-" in code_upper or "/TT-" in code_upper:
        return "Thông tư"
    if "UBTVQH" in code_upper:
        return "Nghị quyết"
    if "QH" in code_upper:
        return "Luật"

    blob = strip_accents(f"{filename} {title}").lower()
    if "bo luat" in blob:
        return "Bộ luật"
    if "nghi dinh" in blob:
        return "Nghị định"
    if "thong tu" in blob:
        return "Thông tư"
    if "nghi quyet" in blob:
        return "Nghị quyết"
    if "vbhn" in blob:
        return "Văn bản hợp nhất"
    if "luat" in blob:
        return "Luật"
    return "Văn bản"


def parse_money(raw: str) -> int:
    """'5.310.000' -> 5310000. Returns 0 when the token is not a number."""

    digits = re.sub(r"[^\d]", "", raw or "")
    return int(digits) if digits else 0


def format_money(value: int) -> str:
    return f"{value:,}".replace(",", ".")


def duration_to_days(quantity: int, unit: str) -> float:
    unit = strip_accents(unit).lower().strip()
    factors = {
        "gio": 1 / 24,
        "ngay": 1.0,
        "ngay lam viec": 1.0,
        "tuan": 7.0,
        "thang": 30.0,
        "nam": 365.0,
    }
    return quantity * factors.get(unit, 1.0)


def canonical_duration(quantity: int, unit: str) -> tuple[str, str]:
    """Return (canonical key, display label) for a duration mention."""

    unit_key = strip_accents(unit).lower().strip()
    labels = {
        "gio": "giờ",
        "ngay": "ngày",
        "ngay lam viec": "ngày làm việc",
        "tuan": "tuần",
        "thang": "tháng",
        "nam": "năm",
    }
    label = labels.get(unit_key, unit_key)
    return f"{quantity}-{slugify(label)}", f"{quantity} {label}"


def detect_issuer(code: str, filename: str) -> str:
    blob = f"{code} {filename}".upper()
    if "TT-BLĐTBXH" in blob or "TT-BLDTBXH" in strip_accents(blob):
        return "Bộ Lao động - Thương binh và Xã hội"
    if "TT-BNV" in blob:
        return "Bộ Nội vụ"
    if "NĐ-CP" in blob or "ND-CP" in strip_accents(blob):
        return "Chính phủ"
    if "UBTVQH" in blob:
        return "Ủy ban Thường vụ Quốc hội"
    if "QH" in blob or "VPQH" in blob:
        return "Quốc hội"
    if "BTC" in blob:
        return "Bộ Tài chính"
    return "Cơ quan nhà nước"


def key_terms(text: str) -> list[str]:
    stop = {
        "theo",
        "quy",
        "dinh",
        "cho",
        "toi",
        "hoi",
        "nhu",
        "nao",
        "ve",
        "va",
        "la",
        "cua",
        "duoc",
        "khong",
        "trong",
        "nhung",
        "gi",
        "cac",
        "mot",
        "so",
    }
    terms: list[str] = []
    for token in VN_WORD_RE.findall(strip_accents(text).lower()):
        if len(token) < 2 or token in stop:
            continue
        terms.append(token)
    return list(dict.fromkeys(terms))


def vector_to_blob(vec: Iterable[float]) -> bytes:
    return array("f", vec).tobytes()


def blob_to_vector(blob: bytes) -> array:
    vec = array("f")
    vec.frombytes(blob)
    return vec


def dot(a: Iterable[float], b: Iterable[float]) -> float:
    return sum(x * y for x, y in zip(a, b))


def _table_lines(table: Any) -> list[str]:
    lines: list[str] = []
    for row in table.rows:
        values: list[str] = []
        seen_cells: set[int] = set()
        for cell in row.cells:
            # Merged cells repeat the same underlying element across the row.
            marker = id(cell._tc)
            if marker in seen_cells:
                continue
            seen_cells.add(marker)
            value = normalize_space(cell.text)
            if value and not is_separator(value):
                values.append(value)
        if values:
            lines.append(" | ".join(values))
    return lines


def _numbering_key(paragraph: Any) -> str:
    """Identify a Word auto-numbering list, if the paragraph belongs to one.

    Word renders the "Điều 3" prefix from list numbering, which never appears in
    ``paragraph.text``. Some laws — Luật An toàn, vệ sinh lao động 84/2015 among
    them — rely on it entirely, so without this the whole document parses as a
    pile of untitled paragraphs and loses every article boundary.
    """

    try:
        properties = paragraph._p.pPr
        numbering = properties.numPr if properties is not None else None
        if numbering is None or numbering.numId is None:
            return "para"
        level = numbering.ilvl.val if numbering.ilvl is not None else 0
        return f"num:{numbering.numId.val}:{level}"
    except Exception:
        return "para"


def docx_blocks(path: Path) -> list[tuple[str, str]]:
    """Read a .docx as ``(kind, text)`` blocks in true document order.

    ``kind`` is ``"para"``, ``"table"`` or ``"num:<listId>:<level>"``. Reading
    paragraphs and tables separately (the previous behaviour) appended every
    table to the end of the document, which detached crucial payloads — the
    regional minimum-wage table of Nghị định 293/2025 for instance — from the
    article that introduces them.
    """

    doc = Document(str(path))
    blocks: list[tuple[str, str]] = []
    try:
        inner = list(doc.iter_inner_content())
    except AttributeError:  # pragma: no cover - python-docx < 1.1
        inner = list(doc.paragraphs) + list(doc.tables)

    for item in inner:
        if hasattr(item, "rows"):
            for line in _table_lines(item):
                blocks.append(("table", line))
            continue
        kind = _numbering_key(item)
        for raw in (getattr(item, "text", "") or "").splitlines():
            text = normalize_space(raw)
            if text and not is_separator(text):
                blocks.append((kind, text))
    return blocks


def docx_lines(path: Path) -> list[str]:
    return [text for _, text in docx_blocks(path)]


class _SyntheticArticle:
    """Stands in for an ARTICLE_RE match when the number came from Word numbering."""

    __slots__ = ("_number", "_heading")

    def __init__(self, number: str, heading: str):
        self._number = number
        self._heading = heading

    def group(self, index: int) -> str:
        return self._number if index == 1 else self._heading


class LegalGraphBuilder:
    def __init__(
        self,
        data_dir: Path,
        storage_dir: Path,
        embedding_config: EmbeddingConfig | None = None,
    ):
        self.data_dir = data_dir
        self.storage_dir = storage_dir
        self.db_path = storage_dir / "legal_graphrag.sqlite"
        self.embedding_config = embedding_config or EmbeddingConfig.from_env()
        self.docs: dict[str, dict[str, Any]] = OrderedDict()
        self.nodes: dict[str, dict[str, Any]] = OrderedDict()
        self.edges: OrderedDict[str, dict[str, Any]] = OrderedDict()
        self.chunks: OrderedDict[str, dict[str, Any]] = OrderedDict()
        self.article_lookup: dict[tuple[str, str], str] = {}
        self.clause_lookup: dict[tuple[str, str, str], str] = {}
        self.point_lookup: dict[tuple[str, str, str, str], str] = {}
        self.doc_guides: dict[str, list[str]] = {}
        self.doc_alias_index: list[tuple[str, str]] = []
        self.mention_budget: dict[str, int] = {}
        self._ascii_cache: dict[str, str] = {}
        self._structural_cache: list[dict[str, Any]] | None = None
        self._article_cache: list[dict[str, Any]] | None = None
        self._children_index: dict[str, list[dict[str, Any]]] | None = None

    def build(self) -> dict[str, int]:
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        paths = sorted(self.data_dir.glob("*.docx"), key=lambda p: slugify(p.name))
        if not paths:
            raise FileNotFoundError(f"No .docx files found in {self.data_dir}")

        for path in paths:
            self._parse_document(path)

        self._finalize_node_text()
        self._register_system_document()

        # Layer 0-1: document lifecycle, cross references, effective dates.
        self._build_document_relations()
        self._build_effective_dates()
        self._build_reference_edges()

        # Layer 2-9: semantic lift over every structural node.
        self._layer2_terms_and_topics()
        self._layer3_wage_and_bonus()
        self._layer4_domain_ontology()
        self._layer5_procedures()
        self._layer6_temporal()
        self._layer7_sanctions_and_risk()
        self._layer8_lifecycles()
        self._layer9_precedents()

        self._refresh_path_labels()
        self._build_chunks()
        self._embed_chunks()
        self._write_sqlite()
        self._write_jsonl()
        return self.summary()

    def summary(self) -> dict[str, Any]:
        node_types: dict[str, int] = {}
        for node in self.nodes.values():
            node_types[node["node_type"]] = node_types.get(node["node_type"], 0) + 1
        relations: dict[str, int] = {}
        for edge in self.edges.values():
            relations[edge["relation"]] = relations.get(edge["relation"], 0) + 1
        return {
            "documents": len(self.docs),
            "nodes": len(self.nodes),
            "edges": len(self.edges),
            "chunks": len(self.chunks),
            "node_types": dict(sorted(node_types.items(), key=lambda kv: -kv[1])),
            "relations": dict(sorted(relations.items(), key=lambda kv: -kv[1])),
        }

    # ------------------------------------------------------------------
    # Shared helpers
    # ------------------------------------------------------------------

    def _ascii(self, text: str) -> str:
        cached = self._ascii_cache.get(text)
        if cached is None:
            # Lexicon patterns are single-line phrases; flatten line breaks so a
            # phrase split across two source paragraphs still matches.
            cached = strip_accents(normalize_space(text)).lower()
            if len(self._ascii_cache) > 40000:
                self._ascii_cache.clear()
            self._ascii_cache[text] = cached
        return cached

    def _semantic_nodes(self) -> list[dict[str, Any]]:
        """Structural nodes, snapshotted once: the layer passes add nodes while
        iterating, and every pass would otherwise rescan the whole graph."""

        if self._structural_cache is None:
            self._structural_cache = [
                node for node in self.nodes.values() if node["node_type"] in {"Điều", "Khoản", "Điểm"}
            ]
        return self._structural_cache

    def _article_nodes(self) -> list[dict[str, Any]]:
        if self._article_cache is None:
            self._article_cache = [node for node in self.nodes.values() if node["node_type"] == "Điều"]
        return self._article_cache

    def _owning_article(self, node: dict[str, Any]) -> str:
        cursor: str | None = node["node_id"]
        seen: set[str] = set()
        while cursor and cursor in self.nodes and cursor not in seen:
            seen.add(cursor)
            current = self.nodes[cursor]
            if current["node_type"] == "Điều":
                return cursor
            cursor = current.get("parent_id")
        return node["node_id"]

    def _concept_node(
        self,
        prefix: str,
        concept: onto.Concept,
        node_type: str,
        doc_id: str = SYSTEM_DOC_ID,
    ) -> str:
        node_id = f"{prefix}:{concept.key}"
        self._add_node(node_id, doc_id, node_type, concept.label, "", concept.label, None, 0)
        if not self.nodes[node_id]["text"]:
            self.nodes[node_id]["text"] = f"{concept.label}. {concept.description}"
        return node_id

    def _link_regulation(self, concept_node_id: str, node: dict[str, Any], relation: str = "QUY_ĐỊNH_TẠI") -> None:
        self._add_edge(concept_node_id, node["node_id"], relation, f"Quy định tại {node['path_label'] or node['label']}")

    def _add_node(
        self,
        node_id: str,
        doc_id: str | None,
        node_type: str,
        label: str,
        number: str = "",
        title: str = "",
        parent_id: str | None = None,
        ordinal: int = 0,
    ) -> str:
        if node_id not in self.nodes:
            self.nodes[node_id] = {
                "node_id": node_id,
                "doc_id": doc_id,
                "node_type": node_type,
                "label": label,
                "number": number,
                "title": title,
                "parent_id": parent_id,
                "path_label": "",
                "text": "",
                "_parts": [],
                "ordinal": ordinal,
            }
        return node_id

    def _append_node_text(self, node_id: str | None, text: str) -> None:
        if node_id and text:
            self.nodes[node_id]["_parts"].append(text)

    def _add_edge(self, source_id: str, target_id: str, relation: str, evidence: str = "") -> None:
        if not source_id or not target_id or source_id == target_id:
            return
        edge_id = hashlib.sha1(f"{source_id}|{relation}|{target_id}|{evidence[:80]}".encode("utf-8")).hexdigest()
        edge_id = f"edge:{edge_id[:20]}"
        if edge_id not in self.edges:
            self.edges[edge_id] = {
                "edge_id": edge_id,
                "source_id": source_id,
                "target_id": target_id,
                "relation": relation,
                "evidence": evidence[:500],
            }

    def _parse_document(self, path: Path) -> None:
        blocks = docx_blocks(path)
        lines = [text for _, text in blocks]
        paragraphs = [text for kind, text in blocks if kind != "table"]
        filename = path.name
        code = detect_code(filename, lines)
        title = smart_doc_title(paragraphs, filename)
        doc_type = detect_doc_type(filename, title, code)
        issuer = detect_issuer(code, filename)
        doc_id = slugify(Path(filename).stem)
        doc_node_id = f"doc:{doc_id}"
        label = f"{title} ({code})" if code else title
        full_text = "\n".join(lines)

        self.docs[doc_id] = {
            "doc_id": doc_id,
            "filename": filename,
            "path": str(path),
            "title": title,
            "code": code,
            "doc_type": doc_type,
            "issuer": issuer,
            "text": full_text,
        }
        self._add_node(doc_node_id, doc_id, "VănBản", label, code, title, None, 0)
        issuer_id = f"issuer:{slugify(issuer)}"
        self._add_node(issuer_id, None, "CơQuanBanHành", issuer, "", issuer, None, 0)
        self._add_edge(issuer_id, doc_node_id, "BAN_HÀNH", issuer)

        kinds = [kind for kind, _ in blocks]
        auto_article_kind = self._auto_article_kind(blocks)
        auto_article_number = 0
        current_chapter: str | None = None
        current_section: str | None = None
        current_article: str | None = None
        current_clause: str | None = None
        current_point: str | None = None
        current_article_number = ""
        intro: list[str] = []
        table_seq = 0
        ordinal = 1
        i = 0
        while i < len(lines):
            text = lines[i]

            if kinds[i] == "table":
                start = i
                while i < len(lines) and kinds[i] == "table":
                    i += 1
                rows = lines[start:i]
                anchor = current_clause or current_article
                table_seq += 1
                ordinal = self._attach_table(
                    doc_id, doc_node_id, anchor, rows, table_seq, ordinal
                )
                if anchor:
                    payload = "\n".join(rows)
                    self._append_node_text(current_article, payload)
                    self._append_node_text(current_clause, payload)
                    self._append_node_text(current_point, payload)
                else:
                    intro.extend(rows)
                continue

            chapter = CHAPTER_RE.match(text)
            if chapter:
                number = chapter.group(1).upper()
                heading = normalize_space(chapter.group(2) or "")
                if not heading and i + 1 < len(lines) and is_heading_title(lines[i + 1]):
                    heading = lines[i + 1]
                    i += 1
                node_id = f"chuong:{doc_id}:{slugify(number)}"
                node_label = f"Chương {number}" + (f". {heading}" if heading else "")
                current_chapter = self._add_node(
                    node_id, doc_id, "Chương", node_label, number, heading, doc_node_id, ordinal
                )
                self._add_edge(current_chapter, doc_node_id, "THUỘC_VỀ", node_label)
                self._append_node_text(current_chapter, node_label)
                current_section = current_article = current_clause = current_point = None
                ordinal += 1
                i += 1
                continue

            section = SECTION_RE.match(text)
            if section:
                number = section.group(1).upper()
                heading = normalize_space(section.group(2) or "")
                if not heading and i + 1 < len(lines) and is_heading_title(lines[i + 1]):
                    heading = lines[i + 1]
                    i += 1
                parent = current_chapter or doc_node_id
                node_id = f"muc:{doc_id}:{slugify(parent)}:{slugify(number)}"
                node_label = f"Mục {number}" + (f". {heading}" if heading else "")
                current_section = self._add_node(
                    node_id, doc_id, "Mục", node_label, number, heading, parent, ordinal
                )
                self._add_edge(current_section, parent, "THUỘC_VỀ", node_label)
                self._append_node_text(current_section, node_label)
                current_article = current_clause = current_point = None
                ordinal += 1
                i += 1
                continue

            article = ARTICLE_RE.match(text)
            if article is None and auto_article_kind and kinds[i] == auto_article_kind:
                auto_article_number += 1
                article = _SyntheticArticle(str(auto_article_number), text)
            if article:
                number = article.group(1)
                heading = normalize_space(article.group(2))
                parent = current_section or current_chapter or doc_node_id
                node_id = f"dieu:{doc_id}:{slugify(number)}"
                node_label = f"Điều {number}. {heading}"
                current_article = self._add_node(
                    node_id, doc_id, "Điều", node_label, number, heading, parent, ordinal
                )
                self.article_lookup[(doc_id, number.lower())] = current_article
                self._add_edge(current_article, parent, "THUỘC_VỀ", node_label)
                self._append_node_text(current_article, node_label)
                current_article_number = number
                current_clause = current_point = None
                ordinal += 1
                i += 1
                continue

            clause = CLAUSE_RE.match(text)
            if clause and current_article:
                number = clause.group(1)
                body = normalize_space(clause.group(2))
                node_id = f"khoan:{doc_id}:{slugify(current_article_number)}:{number}"
                node_label = f"Khoản {number}"
                current_clause = self._add_node(
                    node_id, doc_id, "Khoản", node_label, number, "", current_article, ordinal
                )
                self.clause_lookup[(doc_id, current_article_number.lower(), number)] = current_clause
                self._add_edge(current_clause, current_article, "THUỘC_VỀ", f"{node_label} Điều {current_article_number}")
                line = f"{number}. {body}"
                self._append_node_text(current_article, line)
                self._append_node_text(current_clause, line)
                current_point = None
                ordinal += 1
                i += 1
                continue

            point = POINT_RE.match(text)
            if point and current_article:
                number = point.group(1).lower()
                body = normalize_space(point.group(2))
                parent = current_clause or current_article
                clause_number = self.nodes[parent]["number"] if current_clause else "0"
                node_id = f"diem:{doc_id}:{slugify(current_article_number)}:{slugify(clause_number)}:{slugify(number)}"
                node_label = f"Điểm {number}"
                current_point = self._add_node(
                    node_id, doc_id, "Điểm", node_label, number, "", parent, ordinal
                )
                self.point_lookup[(doc_id, current_article_number.lower(), clause_number, number)] = current_point
                self._add_edge(current_point, parent, "THUỘC_VỀ", f"{node_label} Điều {current_article_number}")
                line = f"{number}) {body}"
                self._append_node_text(current_article, line)
                self._append_node_text(current_clause, line)
                self._append_node_text(current_point, line)
                ordinal += 1
                i += 1
                continue

            if current_article:
                self._append_node_text(current_article, text)
                self._append_node_text(current_clause, text)
                self._append_node_text(current_point, text)
            else:
                intro.append(text)
            i += 1

        if intro:
            self._append_node_text(doc_node_id, "\n".join(intro[:50]))

    @staticmethod
    def _auto_article_kind(blocks: list[tuple[str, str]]) -> str | None:
        """Pick the Word numbering list that stands in for "Điều N" headings.

        Only used when the document has almost no literal ``Điều N.`` lines, so
        normally-formatted documents are untouched.
        """

        explicit = sum(1 for kind, text in blocks if kind != "table" and ARTICLE_RE.match(text))
        if explicit >= 5:
            return None
        counts: dict[str, int] = {}
        for kind, text in blocks:
            if not kind.startswith("num:") or not kind.endswith(":0"):
                continue
            if len(text) > 250 or CLAUSE_RE.match(text) or POINT_RE.match(text):
                continue
            counts[kind] = counts.get(kind, 0) + 1
        if not counts:
            return None
        kind, total = max(counts.items(), key=lambda item: item[1])
        return kind if total >= 5 else None

    def _attach_table(
        self,
        doc_id: str,
        doc_node_id: str,
        anchor_id: str | None,
        rows: list[str],
        table_seq: int,
        ordinal: int,
    ) -> int:
        """Materialise a table as its own node hanging off the enclosing clause.

        Boilerplate tables (letterheads, "Nơi nhận" distribution lists) carry no
        legal payload and are skipped so they never pollute retrieval.
        """

        payload = "\n".join(rows).strip()
        if not payload or len(rows) < 2:
            return ordinal
        ascii_payload = strip_accents(payload).lower()
        boilerplate = ("noi nhan:", "cong hoa xa hoi chu nghia viet nam", "tm. chinh phu", "luu: vt")
        if any(marker in ascii_payload for marker in boilerplate) and "vung" not in ascii_payload:
            return ordinal

        parent = anchor_id or doc_node_id
        node_id = f"bang:{doc_id}:{table_seq}"
        caption = self.nodes[parent]["label"] if parent in self.nodes else "Bảng dữ liệu"
        label = f"Bảng dữ liệu ({caption})"
        self._add_node(node_id, doc_id, "PhụLục_Bảng", label, str(table_seq), caption, parent, ordinal)
        self._append_node_text(node_id, payload)
        self._add_edge(parent, node_id, "CÓ_BẢNG_BIỂU", f"Bảng biểu kèm theo {caption}")
        self._add_edge(node_id, parent, "THUỘC_VỀ", label)
        return ordinal + 1

    def _finalize_node_text(self) -> None:
        for node in self.nodes.values():
            node["text"] = normalize_block("\n".join(node.pop("_parts", [])))
        self._refresh_path_labels()

    def _refresh_path_labels(self) -> None:
        for node_id in list(self.nodes):
            self.nodes[node_id]["path_label"] = self._path_label(node_id)

    def _register_system_document(self) -> None:
        """Synthetic document that owns every cross-corpus semantic entity."""

        self.docs[SYSTEM_DOC_ID] = {
            "doc_id": SYSTEM_DOC_ID,
            "filename": "system",
            "path": "",
            "title": "Bản đồ tri thức LaborCare",
            "code": "SYS",
            "doc_type": "Hệ thống",
            "issuer": "LaborCare",
            "text": (
                "Tập hợp thực thể ngữ nghĩa dùng chung cho toàn bộ đồ thị: thuật ngữ, chủ đề, "
                "khoản thu nhập, thủ tục, chế tài, giai đoạn vòng đời."
            ),
        }

    def _path_label(self, node_id: str) -> str:
        chain: list[str] = []
        cursor = node_id
        seen: set[str] = set()
        while cursor and cursor in self.nodes and cursor not in seen:
            seen.add(cursor)
            node = self.nodes[cursor]
            if node["node_type"] != "CơQuanBanHành":
                chain.append(node["label"])
            cursor = node.get("parent_id")
        return " > ".join(reversed(chain))

    def _doc_aliases(self, doc: dict[str, Any]) -> set[str]:
        """Strings that unambiguously name this document inside another one."""

        title = normalize_space(strip_accents(doc["title"]).lower())
        aliases: set[str] = set()
        code = doc.get("code") or ""
        if code:
            aliases.add(strip_accents(code).lower())
            aliases.add(strip_accents(code.replace("/", "-")).lower())

        type_prefix = strip_accents(doc.get("doc_type") or "").lower()
        if 10 <= len(title) <= 90 and title.startswith(("bo luat", "luat", "nghi dinh", "thong tu", "nghi quyet")):
            aliases.add(title)
        # A bare subject such as "bao hiem xa hoi" appears in every social
        # insurance document, so the short name is only usable when it is
        # qualified by the document type: "luat bao hiem xa hoi".
        compact = re.sub(r"^(bo luat|luat|nghi dinh|thong tu|nghi quyet)\s+", "", title).strip()
        noise = ("quy dinh", "sua doi", "huong dan", "bo sung", "ban hanh", "ve viec")
        if type_prefix and 6 <= len(compact) <= 70 and not compact.startswith(noise):
            aliases.add(f"{type_prefix} {compact}")
        return {alias for alias in aliases if len(alias) >= 8}

    def _build_alias_index(self) -> None:
        """Longest-first alias list used to resolve "Điều X của <văn bản>" refs."""

        entries: list[tuple[str, str]] = []
        for doc_id, doc in self.docs.items():
            if doc_id == SYSTEM_DOC_ID:
                continue
            for alias in self._doc_aliases(doc):
                entries.append((alias, doc_id))
        entries.sort(key=lambda item: len(item[0]), reverse=True)
        self.doc_alias_index = entries

    def _resolve_referenced_doc(self, window_ascii: str, source_doc_id: str) -> str | None:
        """Which document does a citation window point at, if not the current one."""

        if "bo luat nay" in window_ascii or "luat nay" in window_ascii or "nghi dinh nay" in window_ascii:
            return source_doc_id
        for alias, doc_id in self.doc_alias_index:
            if doc_id != source_doc_id and alias in window_ascii:
                return doc_id
        return None

    def _build_document_relations(self) -> None:
        self._build_alias_index()
        doc_aliases = {
            doc_id: self._doc_aliases(doc)
            for doc_id, doc in self.docs.items()
            if doc_id != SYSTEM_DOC_ID
        }
        for source_id, source_doc in self.docs.items():
            if source_id == SYSTEM_DOC_ID:
                continue
            source_text = normalize_space(strip_accents(source_doc["text"][:12000]).lower())
            source_title = normalize_space(strip_accents(source_doc["title"]).lower())
            guide_targets: list[str] = []
            source_type = source_doc["doc_type"]
            source_node = f"doc:{source_id}"
            # An amending document announces itself in its own title
            # ("Luật Sửa đổi, bổ sung một số điều của Luật Bảo hiểm y tế"), so
            # the title is a far cleaner signal than scanning the body.
            is_amendment = bool(re.search(r"sua doi|bo sung", source_title))
            is_replacement = "thay the" in source_title

            for target_id, aliases in doc_aliases.items():
                if source_id == target_id:
                    continue
                target_doc = self.docs[target_id]
                found_alias = next((alias for alias in aliases if alias in source_text), "")
                if not found_alias:
                    continue
                target_node = f"doc:{target_id}"
                named_in_title = any(alias in source_title for alias in aliases)

                if is_amendment and named_in_title:
                    self._add_edge(source_node, target_node, "SỬA_ĐỔI", found_alias)
                if is_replacement and named_in_title:
                    self._add_edge(source_node, target_node, "THAY_THẾ", found_alias)

                is_subordinate = source_type in {"Nghị định", "Thông tư", "Quyết định"}
                is_primary_target = target_doc["doc_type"] in {"Luật", "Bộ luật", "Nghị định"}
                if is_subordinate and is_primary_target and target_doc["doc_type"] != source_type:
                    self._add_edge(source_node, target_node, "HƯỚNG_DẪN", found_alias)
                    guide_targets.append(target_id)
            self.doc_guides[source_id] = guide_targets

    def _build_effective_dates(self) -> None:
        for doc_id, doc in self.docs.items():
            if doc_id == SYSTEM_DOC_ID:
                continue
            match = EFFECTIVE_DATE_RE.search(doc["text"])
            if not match:
                continue
            day, month, year = match.group(1).zfill(2), match.group(2).zfill(2), match.group(3)
            iso = f"{year}-{month}-{day}"
            node_id = f"hieuluc:{iso}"
            label = f"Hiệu lực từ {day}/{month}/{year}"
            self._add_node(node_id, SYSTEM_DOC_ID, "HiệuLựcVănBản", label, iso, label, None, 0)
            self.nodes[node_id]["text"] = f"Mốc thời điểm văn bản có hiệu lực thi hành: ngày {day}/{month}/{year}."
            self._add_edge(f"doc:{doc_id}", node_id, "CÓ_HIỆU_LỰC_TỪ", normalize_space(match.group(0)))
            doc["effective_date"] = iso

    def _build_reference_edges(self) -> None:
        for node_id, node in self.nodes.items():
            if node["node_type"] not in {"Điều", "Khoản", "Điểm"}:
                continue
            text = node.get("text") or ""
            if not text:
                continue
            source_doc_id = node["doc_id"]
            guides = [d for d in self.doc_guides.get(source_doc_id, []) if d != source_doc_id]
            for match in ARTICLE_REF_RE.finditer(text):
                point_no, clause_no, article_no = match.groups()
                article_key = article_no.lower()
                window = text[max(0, match.start() - 110) : match.end() + 140]
                if "Điều này" in window[: 110 + len(match.group(0)) + 12]:
                    continue
                window_ascii = strip_accents(window).lower()

                # An explicit "… của Bộ luật Lao động" wins over the local
                # document; otherwise fall back to self, then to guided laws.
                explicit = self._resolve_referenced_doc(window_ascii, source_doc_id)
                if explicit:
                    target_docs = [explicit]
                else:
                    target_docs = [source_doc_id, *guides]

                for target_doc_id in target_docs:
                    target_id = None
                    if point_no and clause_no:
                        target_id = self.point_lookup.get(
                            (target_doc_id, article_key, clause_no, point_no.lower())
                        )
                    if not target_id and clause_no:
                        target_id = self.clause_lookup.get((target_doc_id, article_key, clause_no))
                    if not target_id:
                        target_id = self.article_lookup.get((target_doc_id, article_key))
                    if target_id:
                        self._add_edge(node_id, target_id, "DẪN_CHIẾU_ĐẾN", normalize_space(window))
                        break

    # ------------------------------------------------------------------
    # Layer 2 - legal terms mined from the corpus + topic map
    # ------------------------------------------------------------------

    def _layer2_terms_and_topics(self) -> None:
        term_nodes = self._mine_definitions()
        self._seed_missing_terms(term_nodes)
        self._link_term_mentions(term_nodes)
        self._build_topic_map()

    def _mine_definitions(self) -> dict[str, str]:
        """Harvest "X là ..." definitions from every "Giải thích từ ngữ" article.

        This replaces the previous hand-written dictionary of 12 terms: every
        law in the corpus carries its own glossary article, so the graph now
        learns hundreds of authoritative definitions straight from the source.
        """

        term_nodes: dict[str, str] = {}
        definition_articles = [
            node for node in self._article_nodes() if DEFINITION_ARTICLE_RE.search(node["title"] or node["label"])
        ]
        for article in definition_articles:
            for clause_id, clause in self._children_of(article["node_id"], "Khoản"):
                body = clause.get("text") or ""
                match = DEFINITION_RE.match(body)
                if not match:
                    continue
                term = normalize_space(match.group("term"))
                term = re.sub(r"^\d{1,3}[\.\)]\s*", "", term).strip()
                if not (3 <= len(term) <= 80) or term.lower().startswith(("trong ", "các từ ngữ")):
                    continue
                definition = normalize_space(match.group("body"))[:900]
                key = slugify(term)
                node_id = f"thuatngu:{key}"
                created = node_id not in self.nodes
                self._add_node(node_id, SYSTEM_DOC_ID, "ThuậtNgữ", term, "", term, None, 0)
                if created or not self.nodes[node_id]["text"]:
                    self.nodes[node_id]["text"] = f"{term} là {definition}"
                self._add_edge(
                    node_id,
                    clause_id,
                    "ĐƯỢC_ĐỊNH_NGHĨA_LÀ",
                    f"Định nghĩa tại {clause['path_label'] or clause['label']}",
                )
                self._add_edge(node_id, article["node_id"], "ĐƯỢC_ĐỊNH_NGHĨA_LÀ", article["label"])
                term_nodes[self._ascii(term)] = node_id
        return term_nodes

    def _children_of(self, parent_id: str, node_type: str) -> list[tuple[str, dict[str, Any]]]:
        return [
            (node_id, node)
            for node_id, node in self.nodes.items()
            if node.get("parent_id") == parent_id and node["node_type"] == node_type
        ]

    def _seed_missing_terms(self, term_nodes: dict[str, str]) -> None:
        for term, definition in onto.SEED_TERMS.items():
            ascii_term = self._ascii(term)
            if ascii_term in term_nodes:
                continue
            node_id = f"thuatngu:{slugify(term)}"
            self._add_node(node_id, SYSTEM_DOC_ID, "ThuậtNgữ", term.title(), "", term.title(), None, 0)
            if not self.nodes[node_id]["text"]:
                self.nodes[node_id]["text"] = f"{term.title()} là {definition}"
            term_nodes[ascii_term] = node_id

    def _link_term_mentions(self, term_nodes: dict[str, str]) -> None:
        """Connect terms to the articles that *use* them.

        Mentions are a weak signal, so they are capped per term and only drawn
        at article level - the old builder emitted 10k DEFINED_AS edges for
        plain mentions, which drowned the real definitions during expansion.
        """

        ranked = sorted(term_nodes.items(), key=lambda kv: -len(kv[0]))
        for ascii_term, node_id in ranked:
            if len(ascii_term) < 8:
                continue
            used = 0
            for article in self._article_nodes():
                if used >= MAX_MENTION_EDGES_PER_TERM:
                    break
                haystack = self._ascii(f"{article['label']} {article.get('text') or ''}")
                if ascii_term not in haystack:
                    continue
                in_title = ascii_term in self._ascii(article["label"])
                if not in_title and haystack.count(ascii_term) < 2:
                    continue
                self._add_edge(node_id, article["node_id"], "ĐỀ_CẬP_ĐẾN", article["label"])
                used += 1

    def _build_topic_map(self) -> None:
        topic_ids: dict[str, str] = {}
        for topic in onto.TOPICS:
            node_id = self._concept_node("chude", topic, "ChủĐề")
            topic_ids[topic.key] = node_id

        doc_topic_hits: dict[str, dict[str, int]] = {}
        for article in self._article_nodes():
            haystack = self._ascii(f"{article['label']} {article.get('text') or ''}")
            scored: list[tuple[int, onto.Concept]] = []
            for topic in onto.TOPICS:
                hits = sum(haystack.count(pattern) for pattern in topic.patterns)
                if hits:
                    scored.append((hits, topic))
            scored.sort(key=lambda item: -item[0])
            for hits, topic in scored[:3]:
                self._add_edge(
                    article["node_id"], topic_ids[topic.key], "THUỘC_CHỦ_ĐỀ", f"{hits} lần đề cập"
                )
                bucket = doc_topic_hits.setdefault(article["doc_id"], {})
                bucket[topic.key] = bucket.get(topic.key, 0) + hits

        for doc_id, hits in doc_topic_hits.items():
            for topic_key, count in sorted(hits.items(), key=lambda kv: -kv[1])[:4]:
                self._add_edge(
                    f"doc:{doc_id}", topic_ids[topic_key], "THUỘC_CHỦ_ĐỀ", f"{count} lần đề cập trong văn bản"
                )

    # ------------------------------------------------------------------
    # Layer 3 - tiền lương & tiền thưởng
    # ------------------------------------------------------------------

    def _layer3_wage_and_bonus(self) -> None:
        component_ids: dict[str, str] = {}
        for concept, parent_key in onto.WAGE_COMPONENTS:
            component_ids[concept.key] = self._concept_node("thunhap", concept, "KhoảnThuNhập")
        for concept, parent_key in onto.WAGE_COMPONENTS:
            if parent_key and parent_key in component_ids:
                self._add_edge(
                    component_ids[concept.key],
                    component_ids[parent_key],
                    "CẤU_THÀNH_LƯƠNG",
                    f"{concept.label} là một cấu phần của {parent_key.replace('-', ' ')}",
                )

        bonus_ids = {c.key: self._concept_node("thuong", c, "LoạiThưởng") for c in onto.BONUS_TYPES}
        form_ids = {c.key: self._concept_node("hinhthuctra", c, "HìnhThứcTrảLương") for c in onto.PAY_FORMS}
        period_ids = {c.key: self._concept_node("kyhantra", c, "KỳHạnTrảLương") for c in onto.PAY_PERIODS}
        base_ids = {c.key: self._concept_node("cancutinh", c, "CănCứTínhLương") for c in onto.WAGE_BASES}
        formula_ids = {c.key: self._concept_node("cachtinh", c, "CáchTính_CôngThức") for c in onto.WAGE_FORMULAS}

        rate_hints = {key: hints for key, hints in onto.WAGE_RATE_HINTS}
        wage_topic = "chude:tien-luong-tien-thuong"

        for node in self._semantic_nodes():
            if node["node_type"] == "Điểm":
                continue
            text = node.get("text") or ""
            if not text:
                continue
            ascii_text = self._ascii(f"{node['label']} {text}")
            if "luong" not in ascii_text and "thuong" not in ascii_text:
                continue

            matched_components: list[str] = []
            for concept, _parent in onto.WAGE_COMPONENTS:
                if concept.matches(ascii_text):
                    matched_components.append(concept.key)
                    self._link_regulation(component_ids[concept.key], node)

            for concept in onto.BONUS_TYPES:
                if concept.matches(ascii_text):
                    self._link_regulation(bonus_ids[concept.key], node)
                    self._add_edge(
                        bonus_ids[concept.key],
                        component_ids["tien-thuong"],
                        "CẤU_THÀNH_LƯƠNG",
                        "Hình thái của tiền thưởng",
                    )

            for concept in onto.PAY_FORMS:
                if concept.matches(ascii_text):
                    self._link_regulation(form_ids[concept.key], node)
                    self._add_edge(
                        component_ids["tien-luong"], form_ids[concept.key], "TRẢ_THEO_HÌNH_THỨC", node["label"]
                    )

            for concept in onto.PAY_PERIODS:
                if concept.matches(ascii_text):
                    self._link_regulation(period_ids[concept.key], node)
                    self._add_edge(
                        component_ids["tien-luong"], period_ids[concept.key], "CÓ_KỲ_HẠN_TRẢ", node["label"]
                    )

            for concept in onto.WAGE_BASES:
                if concept.matches(ascii_text):
                    self._link_regulation(base_ids[concept.key], node)
                    for component_key in matched_components:
                        self._add_edge(
                            component_ids[component_key],
                            base_ids[concept.key],
                            "CĂN_CỨ_TÍNH",
                            f"Căn cứ tính nêu tại {node['label']}",
                        )

            for concept in onto.WAGE_FORMULAS:
                if not concept.matches(ascii_text):
                    continue
                if not any(marker in ascii_text for marker in ("tinh", "muc huong", "bang", "binh quan", "%")):
                    continue
                formula_id = formula_ids[concept.key]
                self._add_edge(formula_id, node["node_id"], "ÁP_DỤNG_CHO", f"Quy định tại {node['label']}")
                self._attach_parameters(formula_id, text, node)

            self._attach_wage_rates(node, text, ascii_text, matched_components, rate_hints, component_ids)

            if "khau tru" in ascii_text and "luong" in ascii_text:
                self._add_edge(
                    component_ids["khau-tru-tien-luong"],
                    component_ids["tien-luong"],
                    "BỊ_KHẤU_TRỪ_TỪ",
                    node["label"],
                )

            if matched_components:
                self._add_edge(node["node_id"], wage_topic, "THUỘC_CHỦ_ĐỀ", "Nội dung về tiền lương/tiền thưởng")

        self._parse_minimum_wage_tables(component_ids)

    def _attach_wage_rates(
        self,
        node: dict[str, Any],
        text: str,
        ascii_text: str,
        matched_components: list[str],
        rate_hints: dict[str, tuple[str, ...]],
        component_ids: dict[str, str],
    ) -> None:
        percents = {match.group(1).replace(",", ".") for match in PERCENT_RE.finditer(text)}
        if not percents:
            return
        for component_key in matched_components:
            hints = rate_hints.get(component_key)
            if not hints or not any(hint in ascii_text for hint in hints):
                continue
            for percent in sorted(percents):
                label = f"{percent}%"
                rate_id = f"tyle:{slugify(label)}"
                self._add_node(rate_id, SYSTEM_DOC_ID, "TỷLệHưởng", label, label, label, None, 0)
                if not self.nodes[rate_id]["text"]:
                    self.nodes[rate_id]["text"] = f"Tỷ lệ luật định {label}."
                self._add_edge(
                    component_ids[component_key],
                    rate_id,
                    "CÓ_MỨC_HƯỞNG",
                    normalize_space(f"{label} theo {node['label']}"),
                )
                self._link_regulation(rate_id, node)

    def _attach_parameters(self, formula_id: str, text: str, node: dict[str, Any]) -> None:
        for match in PERCENT_RE.finditer(text):
            label = f"{match.group(1).replace(',', '.')}%"
            param_id = f"thamso:{slugify(label)}"
            self._add_node(param_id, SYSTEM_DOC_ID, "ThamSố_ConSố", label, label, label, None, 0)
            if not self.nodes[param_id]["text"]:
                self.nodes[param_id]["text"] = f"Tham số tỷ lệ phần trăm luật định: {label}."
            self._add_edge(formula_id, param_id, "CÓ_THAM_SỐ", f"Tỷ lệ {label} tại {node['label']}")
        for match in DURATION_RE.finditer(text):
            quantity = int(match.group(1))
            if quantity > 400:
                continue
            key, label = canonical_duration(quantity, match.group(2))
            param_id = f"thamso:{key}"
            self._add_node(param_id, SYSTEM_DOC_ID, "ThamSố_ConSố", label, str(quantity), label, None, 0)
            if not self.nodes[param_id]["text"]:
                self.nodes[param_id]["text"] = f"Tham số thời gian luật định: {label}."
            self._add_edge(formula_id, param_id, "CÓ_THAM_SỐ", f"Mốc {label} tại {node['label']}")

    def _parse_minimum_wage_tables(self, component_ids: dict[str, str]) -> None:
        """Turn "Vùng I | 5.310.000 | 25.500" rows into queryable wage nodes."""

        for node in list(self.nodes.values()):
            if node["node_type"] != "PhụLục_Bảng":
                continue
            text = node.get("text") or ""
            if "Vùng" not in text:
                continue
            doc = self.docs.get(node["doc_id"], {})
            code = doc.get("code") or node["doc_id"]
            for match in REGION_WAGE_RE.finditer(text):
                region = match.group(1).upper()
                monthly = parse_money(match.group(2))
                hourly = parse_money(match.group(3))
                if monthly < 1_000_000:
                    continue
                node_id = f"luongtoithieu:{slugify(code)}:vung-{region.lower()}"
                label = f"Lương tối thiểu vùng {region} theo {code}"
                self._add_node(node_id, node["doc_id"], "MứcLươngTốiThiểu", label, region, label, node["node_id"], 0)
                self.nodes[node_id]["text"] = (
                    f"Mức lương tối thiểu vùng {region} theo {code}: "
                    f"{format_money(monthly)} đồng/tháng và {format_money(hourly)} đồng/giờ."
                )
                self._add_edge(
                    component_ids["muc-luong-toi-thieu"],
                    node_id,
                    "ÁP_DỤNG_VÙNG",
                    f"Vùng {region}: {format_money(monthly)} đồng/tháng",
                )
                self._add_edge(node_id, node["node_id"], "QUY_ĐỊNH_TẠI", node["label"])
                for amount, unit in ((monthly, "đồng/tháng"), (hourly, "đồng/giờ")):
                    if amount <= 0:
                        continue
                    amount_id = f"sotien:{amount}"
                    amount_label = f"{format_money(amount)} đồng"
                    self._add_node(amount_id, SYSTEM_DOC_ID, "SốTiền", amount_label, str(amount), amount_label, None, 0)
                    if not self.nodes[amount_id]["text"]:
                        self.nodes[amount_id]["text"] = f"Số tiền {amount_label}."
                    self._add_edge(node_id, amount_id, "CÓ_SỐ_TIỀN", f"{amount_label} ({unit})")

    # ------------------------------------------------------------------
    # Layer 4 - subjects, contracts, events, benefits, obligations
    # ------------------------------------------------------------------

    def _layer4_domain_ontology(self) -> None:
        subject_ids = {c.key: self._concept_node("chuthe", c, "ChủThể") for c in onto.SUBJECTS}
        contract_ids = {c.key: self._concept_node("hopdong", c, "HợpĐồngLaoĐộng") for c in onto.CONTRACT_TYPES}
        event_ids = {c.key: self._concept_node("hanhvi", c, "HànhVi_SựKiện") for c in onto.EVENTS}
        benefit_ids = {c.key: self._concept_node("chedo", c, "ChếĐộ_QuyềnLợi") for c in onto.BENEFITS}
        obligation_ids = {c.key: self._concept_node("nghiavu", c, "NghĩaVụ") for c in onto.OBLIGATIONS}

        prohibition_markers = ("nghiem cam", "khong duoc", "bi cam", "cam nguoi su dung lao dong", "trai phap luat")

        for node in self._semantic_nodes():
            if node["node_type"] == "Điểm":
                continue
            text = node.get("text") or ""
            if not text:
                continue
            ascii_text = self._ascii(f"{node['label']} {text}")

            active_subjects = [
                key for key, concept in ((c.key, c) for c in onto.SUBJECTS) if concept.matches(ascii_text)
            ]
            for key in active_subjects:
                self._link_regulation(subject_ids[key], node)

            active_contracts = [c.key for c in onto.CONTRACT_TYPES if c.matches(ascii_text)]
            for key in active_contracts:
                self._link_regulation(contract_ids[key], node)
                for subject_key in active_subjects:
                    if subject_key in {"nguoi-lao-dong", "nguoi-su-dung-lao-dong"}:
                        self._add_edge(
                            subject_ids[subject_key], contract_ids[key], "KÝ_KẾT", node["label"]
                        )

            is_prohibition = any(marker in ascii_text for marker in prohibition_markers)
            for concept in onto.EVENTS:
                if not concept.matches(ascii_text):
                    continue
                event_id = event_ids[concept.key]
                self._link_regulation(event_id, node)
                for subject_key in active_subjects:
                    self._add_edge(subject_ids[subject_key], event_id, "THỰC_HIỆN", node["label"])
                if is_prohibition:
                    self._add_edge(
                        event_id, node["node_id"], "BỊ_NGHIÊM_CẤM", normalize_space(text[:220])
                    )

            for concept in onto.BENEFITS:
                if not concept.matches(ascii_text):
                    continue
                benefit_id = benefit_ids[concept.key]
                self._link_regulation(benefit_id, node)
                self._add_edge(
                    subject_ids["nguoi-lao-dong"], benefit_id, "CÓ_QUYỀN_HƯỞNG", node["label"]
                )

            for concept in onto.OBLIGATIONS:
                if not concept.matches(ascii_text):
                    continue
                obligation_id = obligation_ids[concept.key]
                self._link_regulation(obligation_id, node)
                holder = "nguoi-su-dung-lao-dong" if "nguoi su dung lao dong" in ascii_text else None
                if holder:
                    self._add_edge(subject_ids[holder], obligation_id, "CÓ_NGHĨA_VỤ", node["label"])

    # ------------------------------------------------------------------
    # Layer 5 - administrative procedures
    # ------------------------------------------------------------------

    def _layer5_procedures(self) -> None:
        procedure_ids = {c.key: self._concept_node("thutuc", c, "ThủTục_ChếĐộ") for c in onto.PROCEDURES}
        dossier_ids = {c.key: self._concept_node("hoso", c, "HồSơ_GiấyTờ") for c in onto.DOSSIERS}
        condition_ids = {c.key: self._concept_node("dieukien", c, "ĐiềuKiện") for c in onto.CONDITIONS}
        agency_ids = {c.key: self._concept_node("coquan", c, "CơQuanGiảiQuyết") for c in onto.AGENCIES}

        for node in self._semantic_nodes():
            if node["node_type"] == "Điểm":
                continue
            text = node.get("text") or ""
            if not text:
                continue
            ascii_text = self._ascii(f"{node['label']} {text}")

            active_procedures = [c.key for c in onto.PROCEDURES if c.matches(ascii_text)]
            if not active_procedures:
                continue
            for key in active_procedures:
                self._link_regulation(procedure_ids[key], node)

            for concept in onto.DOSSIERS:
                if concept.matches(ascii_text):
                    self._link_regulation(dossier_ids[concept.key], node)
                    for key in active_procedures:
                        self._add_edge(
                            procedure_ids[key], dossier_ids[concept.key], "BAO_GỒM_HỒ_SƠ", node["label"]
                        )

            for concept in onto.CONDITIONS:
                if concept.matches(ascii_text):
                    self._link_regulation(condition_ids[concept.key], node)
                    for key in active_procedures:
                        self._add_edge(
                            procedure_ids[key], condition_ids[concept.key], "YÊU_CẦU_ĐIỀU_KIỆN", node["label"]
                        )

            for concept in onto.AGENCIES:
                if concept.matches(ascii_text):
                    self._link_regulation(agency_ids[concept.key], node)
                    for key in active_procedures:
                        self._add_edge(procedure_ids[key], agency_ids[concept.key], "NỘP_TẠI", node["label"])

            if any(marker in ascii_text for marker in ("thoi han", "trong thoi gian", "ke tu ngay nhan du ho so")):
                for match in DURATION_RE.finditer(text):
                    quantity = int(match.group(1))
                    if quantity > 120:
                        continue
                    key, label = canonical_duration(quantity, match.group(2))
                    deadline_id = f"thoihan:{key}"
                    self._add_node(
                        deadline_id, SYSTEM_DOC_ID, "ThờiHạn_ThủTục", label, str(quantity), label, None, 0
                    )
                    if not self.nodes[deadline_id]["text"]:
                        self.nodes[deadline_id]["text"] = f"Thời hạn xử lý thủ tục: {label}."
                    for procedure_key in active_procedures:
                        self._add_edge(
                            procedure_ids[procedure_key],
                            deadline_id,
                            "CÓ_THỜI_HẠN_LÀ",
                            normalize_space(f"{label} theo {node['label']}"),
                        )

    # ------------------------------------------------------------------
    # Layer 6 - temporal reasoning
    # ------------------------------------------------------------------

    def _layer6_temporal(self) -> None:
        trigger_ids = {c.key: self._concept_node("sukien", c, "SựKiệnKíchHoạt") for c in onto.TIME_TRIGGERS}
        state_ids = {c.key: self._concept_node("trangthai", c, "TrạngTháiPhápLý") for c in onto.LEGAL_STATES}

        for node in self._semantic_nodes():
            if node["node_type"] == "Điểm":
                continue
            text = node.get("text") or ""
            if not text:
                continue
            ascii_text = self._ascii(f"{node['label']} {text}")
            if not any(marker in ascii_text for marker in ("ke tu", "thoi hieu", "thoi han", "trong thoi han")):
                continue

            active_triggers = [c.key for c in onto.TIME_TRIGGERS if c.matches(ascii_text)]
            if not active_triggers:
                continue
            for key in active_triggers:
                self._link_regulation(trigger_ids[key], node)

            durations: list[tuple[str, str]] = []
            for match in DURATION_RE.finditer(text):
                quantity = int(match.group(1))
                if quantity > 400:
                    continue
                durations.append(canonical_duration(quantity, match.group(2)))
            for key, label in durations[:6]:
                milestone_id = f"mocthoigian:{key}"
                self._add_node(
                    milestone_id, SYSTEM_DOC_ID, "MốcThờiGian_LuậtĐịnh", label, key, label, None, 0
                )
                if not self.nodes[milestone_id]["text"]:
                    self.nodes[milestone_id]["text"] = f"Mốc thời gian luật định: {label}."
                self._link_regulation(milestone_id, node)
                for trigger_key in active_triggers:
                    self._add_edge(
                        trigger_ids[trigger_key],
                        milestone_id,
                        "BẮT_ĐẦU_TÍNH_THỜI_HIỆU",
                        normalize_space(f"{label} kể từ {trigger_key.replace('sk-', '').replace('-', ' ')}"),
                    )
                for concept in onto.LEGAL_STATES:
                    if concept.matches(ascii_text):
                        self._add_edge(
                            milestone_id, state_ids[concept.key], "CHUYỂN_TRẠNG_THÁI", node["label"]
                        )
                        self._link_regulation(state_ids[concept.key], node)

    # ------------------------------------------------------------------
    # Layer 7 - sanctions, fines and risk
    # ------------------------------------------------------------------

    def _layer7_sanctions_and_risk(self) -> None:
        risk_ids: dict[str, str] = {}
        for key, label, description in onto.RISK_LEVELS:
            node_id = f"ruiro:{key}"
            self._add_node(node_id, SYSTEM_DOC_ID, "MứcĐộRủiRo", label, key, label, None, 0)
            self.nodes[node_id]["text"] = f"{label}. {description}"
            risk_ids[key] = node_id

        violation_ids = {c.key: self._concept_node("vipham", c, "HànhViViPhạm") for c in onto.VIOLATIONS}
        remedy_ids = {c.key: self._concept_node("khacphuc", c, "BiệnPhápKhắcPhục") for c in onto.REMEDIES}
        extra_ids = {c.key: self._concept_node("xpbosung", c, "HìnhThứcXửPhạtBổSung") for c in onto.EXTRA_SANCTIONS}

        self._mine_penalty_articles(risk_ids, remedy_ids, extra_ids)

        for node in self._semantic_nodes():
            text = node.get("text") or ""
            if not text:
                continue
            ascii_text = self._ascii(f"{node['label']} {text}")
            fines = self._extract_fines(text)
            active_violations = [c.key for c in onto.VIOLATIONS if c.matches(ascii_text)]
            has_sanction_context = bool(fines) or any(
                marker in ascii_text
                for marker in ("nghiem cam", "vi pham", "xu phat", "bien phap khac phuc hau qua")
            )
            if not active_violations or not has_sanction_context:
                continue

            for key in active_violations:
                violation_id = violation_ids[key]
                self._link_regulation(violation_id, node)

                for low, high in fines:
                    fine_id = f"phat:{low}-{high}"
                    label = f"Phạt tiền từ {format_money(low)} đến {format_money(high)} đồng"
                    self._add_node(fine_id, SYSTEM_DOC_ID, "MứcPhạtTiền", label, str(low), label, None, 0)
                    if not self.nodes[fine_id]["text"]:
                        self.nodes[fine_id]["text"] = (
                            f"{label}. Mức phạt với tổ chức thường bằng 02 lần mức phạt với cá nhân."
                        )
                    self._add_edge(violation_id, fine_id, "BỊ_XỬ_PHẠT", normalize_space(node["label"]))
                    self._link_regulation(fine_id, node)
                    self._add_edge(violation_id, risk_ids[self._risk_bucket(high, ascii_text)], "GÂY_RA_RỦI_RO", label)

                if not fines:
                    self._add_edge(
                        violation_id, risk_ids[self._risk_bucket(0, ascii_text)], "GÂY_RA_RỦI_RO", node["label"]
                    )

                for concept in onto.REMEDIES:
                    if concept.matches(ascii_text):
                        self._add_edge(violation_id, remedy_ids[concept.key], "KHẮC_PHỤC_BẰNG", node["label"])
                        self._link_regulation(remedy_ids[concept.key], node)

                for concept in onto.EXTRA_SANCTIONS:
                    if concept.matches(ascii_text):
                        self._add_edge(
                            violation_id, extra_ids[concept.key], "BỊ_XỬ_PHẠT_BỔ_SUNG", node["label"]
                        )
                        self._link_regulation(extra_ids[concept.key], node)

    def _mine_penalty_articles(
        self,
        risk_ids: dict[str, str],
        remedy_ids: dict[str, str],
        extra_ids: dict[str, str],
    ) -> None:
        """Read the sanction decrees end to end.

        Penalty articles name their own offence in the heading — "Điều 17. Vi
        phạm quy định về tiền lương" — and place each fine bracket in a clause
        or point beneath. Mining those headings covers every offence in Nghị
        định 12/2022 instead of the handful the curated lexicon anticipated.
        """

        for article in self._article_nodes():
            heading = article["title"] or article["label"]
            match = re.match(
                r"^(?:Vi phạm (?:quy định|các quy định)?\s*(?:về|đối với|trong)?|Hành vi vi phạm)\s+(.{4,140})$",
                heading,
                re.IGNORECASE,
            )
            if not match:
                continue
            subject = normalize_space(match.group(1)).rstrip(".")
            doc = self.docs.get(article["doc_id"], {})
            code = doc.get("code") or article["doc_id"]
            violation_id = f"vipham:{slugify(article['doc_id'])}:{slugify(subject)[:60]}"
            label = f"Vi phạm quy định về {subject}"
            self._add_node(
                violation_id, article["doc_id"], "HànhViViPhạm", label, "", label, article["node_id"], 0
            )
            if not self.nodes[violation_id]["text"]:
                self.nodes[violation_id]["text"] = f"{label} — chế tài quy định tại {article['label']} ({code})."
            self._link_regulation(violation_id, article)

            # Attach every fine bracket declared under this article.
            for descendant in self._descendants_of(article["node_id"]):
                text = descendant.get("text") or ""
                fines = self._extract_fines(text)
                if not fines:
                    continue
                ascii_text = self._ascii(text)
                for low, high in fines:
                    fine_id = f"phat:{low}-{high}"
                    fine_label = f"Phạt tiền từ {format_money(low)} đến {format_money(high)} đồng"
                    self._add_node(fine_id, SYSTEM_DOC_ID, "MứcPhạtTiền", fine_label, str(low), fine_label, None, 0)
                    if not self.nodes[fine_id]["text"]:
                        self.nodes[fine_id]["text"] = (
                            f"{fine_label}. Mức phạt tiền với tổ chức thường bằng 02 lần mức phạt với cá nhân."
                        )
                    self._add_edge(
                        violation_id,
                        fine_id,
                        "BỊ_XỬ_PHẠT",
                        normalize_space(f"{fine_label} tại {descendant['path_label'] or descendant['label']}"),
                    )
                    self._link_regulation(fine_id, descendant)
                    self._add_edge(
                        violation_id,
                        risk_ids[self._risk_bucket(high, ascii_text)],
                        "GÂY_RA_RỦI_RO",
                        fine_label,
                    )
                for concept in onto.REMEDIES:
                    if concept.matches(ascii_text):
                        self._add_edge(violation_id, remedy_ids[concept.key], "KHẮC_PHỤC_BẰNG", descendant["label"])
                        self._link_regulation(remedy_ids[concept.key], descendant)
                for concept in onto.EXTRA_SANCTIONS:
                    if concept.matches(ascii_text):
                        self._add_edge(
                            violation_id, extra_ids[concept.key], "BỊ_XỬ_PHẠT_BỔ_SUNG", descendant["label"]
                        )

    def _descendants_of(self, root_id: str) -> list[dict[str, Any]]:
        if self._children_index is None:
            index: dict[str, list[dict[str, Any]]] = {}
            for node in self.nodes.values():
                parent = node.get("parent_id")
                if parent:
                    index.setdefault(parent, []).append(node)
            self._children_index = index
        collected: list[dict[str, Any]] = []
        stack = list(self._children_index.get(root_id, []))
        while stack:
            node = stack.pop()
            collected.append(node)
            stack.extend(self._children_index.get(node["node_id"], []))
        return collected

    @staticmethod
    def _extract_fines(text: str) -> list[tuple[int, int]]:
        ranges: list[tuple[int, int]] = []
        for match in FINE_RANGE_RE.finditer(text):
            low, high = parse_money(match.group(1)), parse_money(match.group(2))
            if low and high and low <= high:
                ranges.append((low, high))
        return ranges[:6]

    @staticmethod
    def _risk_bucket(ceiling: int, ascii_text: str) -> str:
        if any(marker in ascii_text for marker in ("truy cuu trach nhiem hinh su", "dinh chi hoat dong", "truc xuat")):
            return "nghiem-trong"
        for threshold, key in onto.RISK_THRESHOLDS:
            if ceiling and ceiling < threshold:
                return key
        return "cao" if ceiling else "thap"

    # ------------------------------------------------------------------
    # Layer 8 - lifecycles
    # ------------------------------------------------------------------

    def _layer8_lifecycles(self) -> None:
        for prefix, node_type, stages in (
            ("lifecycle_nld", "GiaiĐoạn_NLĐ", onto.LIFECYCLE_NLD),
            ("lifecycle_dn", "GiaiĐoạn_DoanhNghiệp", onto.LIFECYCLE_DN),
        ):
            stage_ids: list[str] = []
            for key, label, description, _patterns in stages:
                node_id = f"{prefix}:{key}"
                self._add_node(node_id, SYSTEM_DOC_ID, node_type, label, "", label, None, 0)
                self.nodes[node_id]["text"] = f"{label}. {description}"
                stage_ids.append(node_id)
            for previous, following in zip(stage_ids, stage_ids[1:]):
                self._add_edge(
                    previous,
                    following,
                    "GIAI_ĐOẠN_TIẾP_THEO",
                    f"{self.nodes[previous]['label']} → {self.nodes[following]['label']}",
                )

            budgets: dict[str, int] = {}
            for article in self._article_nodes():
                haystack = self._ascii(f"{article['label']} {article.get('text') or ''}")
                for key, label, _description, patterns in stages:
                    node_id = f"{prefix}:{key}"
                    if budgets.get(key, 0) >= 160:
                        continue
                    if any(pattern in haystack for pattern in patterns):
                        self._add_edge(
                            node_id, article["node_id"], "KÍCH_HOẠT_NGHĨA_VỤ", article["label"]
                        )
                        budgets[key] = budgets.get(key, 0) + 1

    # ------------------------------------------------------------------
    # Layer 9 - precedents (populated when a case-law corpus is added)
    # ------------------------------------------------------------------

    def _layer9_precedents(self) -> None:
        facts = (
            ("tranh chap hoc nghe", "Tranh chấp về hợp đồng học nghề"),
            ("sa thai trai phap luat", "Sa thải người lao động trái luật"),
            ("don phuong cham dut hop dong lao dong trai phap luat", "Đơn phương chấm dứt HĐLĐ trái pháp luật"),
            ("no luong", "Doanh nghiệp chậm trả, nợ lương người lao động"),
        )
        pattern = re.compile(r"(Án lệ số \d+/20\d\d/AL|Án lệ số \d+|Bản án số \d+/20\d\d/[A-ZĐ0-9\-]+)", re.I)
        for node in self._semantic_nodes():
            text = node.get("text") or ""
            if "án lệ" not in text.lower() and "bản án số" not in text.lower():
                continue
            ascii_text = self._ascii(text)
            for match in pattern.finditer(text):
                reference = normalize_space(match.group(0))
                case_id = f"anle:{slugify(reference)}"
                self._add_node(case_id, SYSTEM_DOC_ID, "ÁnLệ", reference, "", reference, None, 0)
                if not self.nodes[case_id]["text"]:
                    self.nodes[case_id]["text"] = f"Án lệ/bản án được viện dẫn: {reference}."
                self._add_edge(case_id, node["node_id"], "ÁP_DỤNG_ĐIỀU_LUẬT", node["label"])
                for keyword, label in facts:
                    if keyword not in ascii_text:
                        continue
                    fact_id = f"tinhtiet:{slugify(keyword)}"
                    ruling_id = f"phanquyet:{slugify(keyword)}"
                    self._add_node(fact_id, SYSTEM_DOC_ID, "TìnhTiếtCốtLõi", label, "", label, None, 0)
                    self.nodes[fact_id]["text"] = f"Tình tiết cốt lõi: {label}."
                    self._add_node(
                        ruling_id, SYSTEM_DOC_ID, "PhánQuyết", f"Phán quyết về {label.lower()}", "", label, None, 0
                    )
                    self.nodes[ruling_id]["text"] = f"Hướng phán quyết của toà án đối với tình tiết: {label}."
                    self._add_edge(case_id, fact_id, "CÓ_TÌNH_TIẾT_TƯƠNG_TỰ", label)
                    self._add_edge(fact_id, ruling_id, "DẪN_ĐẾN_PHÁN_QUYẾT", label)


    def _add_chunk(
        self,
        doc_id: str,
        node_id: str,
        chunk_type: str,
        text: str,
        ordinal: int,
        title: str = "",
    ) -> None:
        text = normalize_block(text)
        if not text or token_count(text) < 4:
            return
        node = self.nodes[node_id]
        citation = node["path_label"]
        chunk_id = f"chunk:{doc_id}:{slugify(node_id)}:{chunk_type}:{ordinal}"
        if chunk_id in self.chunks:
            return
        heading = title or node["label"]
        self.chunks[chunk_id] = {
            "chunk_id": chunk_id,
            "doc_id": doc_id,
            "node_id": node_id,
            "chunk_type": chunk_type,
            "title": heading,
            "path_label": citation,
            "citation": citation,
            "text": text,
            "token_count": token_count(text),
            "ordinal": ordinal,
            "vector": b"",
        }

    def _add_windowed_chunks(
        self,
        doc_id: str,
        node_id: str,
        primary_type: str,
        text: str,
        ordinal: int,
    ) -> int:
        for index, window in enumerate(embedding_text_windows(text)):
            self._add_chunk(
                doc_id,
                node_id,
                primary_type if index == 0 else "sliding",
                window,
                ordinal,
            )
            ordinal += 1
        return ordinal

    def _embed_chunks(self) -> None:
        rows = list(self.chunks.values())
        texts = [f"{row['title']}\n{row['path_label']}\n{row['text']}" for row in rows]
        service = get_embedding_service(self.embedding_config)
        checkpoint_enabled = os.getenv(
            "LEGAL_EMBEDDING_CHECKPOINT_ENABLED",
            "",
        ).strip().lower() in {"1", "true", "yes", "on"}
        checkpoint: PostgresEmbeddingCheckpoint | None = None
        cached: dict[str, tuple[str, list[float]]] = {}
        if checkpoint_enabled:
            checkpoint = PostgresEmbeddingCheckpoint(
                os.getenv("DATABASE_URL", ""),
                self.embedding_config,
            )
            cached = checkpoint.load()

        pending: list[tuple[dict[str, Any], str, str]] = []
        restored = 0
        for row, text in zip(rows, texts, strict=True):
            content_hash = embedding_content_hash(text)
            cached_row = cached.get(row["chunk_id"])
            if cached_row and cached_row[0] == content_hash:
                row["vector"] = vector_to_blob(cached_row[1])
                restored += 1
            else:
                pending.append((row, text, content_hash))

        if checkpoint_enabled:
            print(
                "Embedding checkpoint: "
                f"restored {restored}/{len(rows)}; pending {len(pending)}.",
                flush=True,
            )

        checkpoint_batch_size = len(pending) or 1
        if checkpoint_enabled:
            checkpoint_batch_size = max(
                1,
                int(
                    os.getenv(
                        "LEGAL_EMBEDDING_CHECKPOINT_BATCH_SIZE",
                        "640",
                    )
                ),
            )
        completed = restored
        for offset in range(0, len(pending), checkpoint_batch_size):
            batch = pending[offset : offset + checkpoint_batch_size]
            embeddings = service.embed_documents(
                [text for _, text, _ in batch],
                show_progress=True,
            )
            records: list[EmbeddingCheckpointRecord] = []
            for (row, _text, content_hash), embedding in zip(
                batch,
                embeddings,
                strict=True,
            ):
                row["vector"] = vector_to_blob(embedding)
                records.append(
                    EmbeddingCheckpointRecord(
                        chunk_id=row["chunk_id"],
                        content_hash=content_hash,
                        vector=embedding,
                    )
                )
            if checkpoint is not None:
                checkpoint.save(records)
            completed += len(batch)
            if checkpoint_enabled:
                print(
                    f"Embedding checkpoint: saved {completed}/{len(rows)}.",
                    flush=True,
                )

    def _semantic_chunk_text(self, node_id: str, node: dict[str, Any], outgoing: dict[str, list[str]]) -> str:
        """Give a concept node a body worth embedding.

        A bare "Tiền lương. Số tiền …" line is too thin to win a vector search.
        Folding in the citations the concept points at turns each concept node
        into a navigable hub chunk: the retriever can land on "Tiền lương" and
        immediately see Điều 90, 95, 97, 104 …
        """

        base = node.get("text") or node["label"]
        citations = outgoing.get(node_id) or []
        if not citations:
            return base
        listed = "; ".join(dict.fromkeys(citations))[:1200]
        return f"{base} Căn cứ pháp lý liên quan: {listed}."

    def _build_chunks(self) -> None:
        relevant = {"QUY_ĐỊNH_TẠI", "ĐƯỢC_ĐỊNH_NGHĨA_LÀ", "ÁP_DỤNG_CHO", "BỊ_NGHIÊM_CẤM", "KÍCH_HOẠT_NGHĨA_VỤ"}
        outgoing: dict[str, list[str]] = {}
        for edge in self.edges.values():
            if edge["relation"] not in relevant:
                continue
            target = self.nodes.get(edge["target_id"])
            if not target:
                continue
            label = target.get("path_label") or target["label"]
            bucket = outgoing.setdefault(edge["source_id"], [])
            if len(bucket) < 14:
                bucket.append(label)

        ordinal = 0
        for node_id, node in self.nodes.items():
            doc_id = node["doc_id"]
            if not doc_id:
                continue
            node_type = node["node_type"]
            text = node.get("text") or node["label"]
            if node_type in {"Chương", "Mục"}:
                self._add_chunk(doc_id, node_id, "structure", text, ordinal)
                ordinal += 1
            elif node_type == "VănBản":
                ordinal = self._add_windowed_chunks(
                    doc_id,
                    node_id,
                    "document_intro",
                    text,
                    ordinal,
                )
            elif node_type == "Điều":
                ordinal = self._add_windowed_chunks(
                    doc_id,
                    node_id,
                    "article",
                    text,
                    ordinal,
                )
            elif node_type == "Khoản":
                ordinal = self._add_windowed_chunks(
                    doc_id,
                    node_id,
                    "clause",
                    text,
                    ordinal,
                )
            elif node_type == "Điểm":
                ordinal = self._add_windowed_chunks(
                    doc_id,
                    node_id,
                    "point",
                    text,
                    ordinal,
                )
            elif node_type == "PhụLục_Bảng":
                ordinal = self._add_windowed_chunks(
                    doc_id,
                    node_id,
                    "table",
                    text,
                    ordinal,
                )
            else:
                self._add_chunk(
                    doc_id, node_id, "semantic", self._semantic_chunk_text(node_id, node, outgoing), ordinal
                )
                ordinal += 1

    def _write_sqlite(self) -> None:
        if self.db_path.exists():
            self.db_path.unlink()
        conn = sqlite3.connect(self.db_path)
        conn.execute("PRAGMA journal_mode=WAL")
        conn.executescript(
            """
            CREATE TABLE docs (
                doc_id TEXT PRIMARY KEY,
                filename TEXT,
                path TEXT,
                title TEXT,
                code TEXT,
                doc_type TEXT,
                issuer TEXT,
                text TEXT
            );
            CREATE TABLE nodes (
                node_id TEXT PRIMARY KEY,
                doc_id TEXT,
                node_type TEXT,
                label TEXT,
                number TEXT,
                title TEXT,
                parent_id TEXT,
                path_label TEXT,
                text TEXT,
                ordinal INTEGER
            );
            CREATE TABLE edges (
                edge_id TEXT PRIMARY KEY,
                source_id TEXT,
                target_id TEXT,
                relation TEXT,
                evidence TEXT
            );
            CREATE TABLE chunks (
                chunk_id TEXT PRIMARY KEY,
                doc_id TEXT,
                node_id TEXT,
                chunk_type TEXT,
                title TEXT,
                path_label TEXT,
                citation TEXT,
                text TEXT,
                token_count INTEGER,
                ordinal INTEGER,
                vector BLOB
            );
            CREATE TABLE index_metadata (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL
            );
            CREATE VIRTUAL TABLE chunk_fts USING fts5(
                chunk_id UNINDEXED,
                title,
                path_label,
                citation,
                text,
                tokenize='unicode61 remove_diacritics 2'
            );
            CREATE INDEX idx_nodes_parent ON nodes(parent_id);
            CREATE INDEX idx_nodes_doc_type ON nodes(doc_id, node_type);
            CREATE INDEX idx_edges_source ON edges(source_id, relation);
            CREATE INDEX idx_edges_target ON edges(target_id, relation);
            CREATE INDEX idx_chunks_node ON chunks(node_id);
            CREATE INDEX idx_chunks_doc ON chunks(doc_id);
            """
        )
        conn.executemany(
            "INSERT INTO docs VALUES (:doc_id, :filename, :path, :title, :code, :doc_type, :issuer, :text)",
            self.docs.values(),
        )
        node_rows = [
            {k: v for k, v in node.items() if k in {"node_id", "doc_id", "node_type", "label", "number", "title", "parent_id", "path_label", "text", "ordinal"}}
            for node in self.nodes.values()
        ]
        conn.executemany(
            "INSERT INTO nodes VALUES (:node_id, :doc_id, :node_type, :label, :number, :title, :parent_id, :path_label, :text, :ordinal)",
            node_rows,
        )
        conn.executemany(
            "INSERT INTO edges VALUES (:edge_id, :source_id, :target_id, :relation, :evidence)",
            self.edges.values(),
        )
        conn.executemany(
            "INSERT INTO chunks VALUES (:chunk_id, :doc_id, :node_id, :chunk_type, :title, :path_label, :citation, :text, :token_count, :ordinal, :vector)",
            self.chunks.values(),
        )
        conn.executemany(
            "INSERT INTO index_metadata(key, value) VALUES (?, ?)",
            [
                ("embedding_model", self.embedding_config.model),
                ("embedding_revision", self.embedding_config.model_revision),
                ("embedding_dimensions", str(self.embedding_config.dimensions)),
            ],
        )
        conn.executemany(
            "INSERT INTO chunk_fts(chunk_id, title, path_label, citation, text) VALUES (:chunk_id, :title, :path_label, :citation, :text)",
            self.chunks.values(),
        )
        conn.commit()
        conn.close()

    def _write_jsonl(self) -> None:
        exports = {
            "documents.jsonl": self.docs.values(),
            "nodes.jsonl": (
                {k: v for k, v in node.items() if k != "vector"} for node in self.nodes.values()
            ),
            "edges.jsonl": self.edges.values(),
            "chunks.jsonl": (
                {k: v for k, v in chunk.items() if k != "vector"} for chunk in self.chunks.values()
            ),
        }
        for filename, rows in exports.items():
            with (self.storage_dir / filename).open("w", encoding="utf-8") as f:
                for row in rows:
                    f.write(json.dumps(row, ensure_ascii=False) + "\n")


class GraphRAGStore:
    def __init__(
        self,
        db_path: Path | str | None = None,
        embedding_config: EmbeddingConfig | None = None,
    ):
        self.db_path = Path(db_path or os.getenv("LEGAL_GRAPHRAG_DB", DEFAULT_DB_PATH))
        self.embedding_config = embedding_config or EmbeddingConfig.from_env()
        if not self.db_path.exists():
            raise FileNotFoundError(f"GraphRAG index not found: {self.db_path}")
        self.conn = sqlite3.connect(self.db_path, check_same_thread=False)
        self.conn.row_factory = sqlite3.Row
        self._vectors: list[tuple[str, bytes]] | None = None
        self._matrix: Any = None
        self._matrix_ids: list[str] = []
        self._terms: list[tuple[str, str]] | None = None
        self._doc_bridges: dict[str, list[str]] | None = None
        self._node_chunk_cache: dict[str, sqlite3.Row | None] = {}
        self._doc_types: dict[str, tuple[str, str]] | None = None
        self._concepts: list[tuple[str, str]] | None = None
        try:
            self._validate_embedding_metadata()
        except Exception:
            self.conn.close()
            raise

    def _validate_embedding_metadata(self) -> None:
        try:
            rows = self.conn.execute("SELECT key, value FROM index_metadata").fetchall()
        except sqlite3.OperationalError as exc:
            raise RuntimeError(
                "Local GraphRAG index uses legacy vectors; rebuild it with the configured Vertex AI embedding model."
            ) from exc
        metadata = {row["key"]: row["value"] for row in rows}
        expected = {
            "embedding_model": self.embedding_config.model,
            "embedding_revision": self.embedding_config.model_revision,
            "embedding_dimensions": str(self.embedding_config.dimensions),
        }
        if any(metadata.get(key) != value for key, value in expected.items()):
            raise RuntimeError(
                f"Local GraphRAG embedding metadata {metadata!r} does not match {expected!r}; rebuild the index."
            )

    def close(self) -> None:
        self.conn.close()

    def stats(self) -> dict[str, Any]:
        def count(table: str) -> int:
            return int(self.conn.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0])

        relation_rows = self.conn.execute(
            "SELECT relation, COUNT(*) AS count FROM edges GROUP BY relation ORDER BY count DESC"
        ).fetchall()
        node_types_rows = self.conn.execute(
            "SELECT node_type, COUNT(*) AS count FROM nodes GROUP BY node_type ORDER BY count DESC"
        ).fetchall()
        return {
            "documents": count("docs"),
            "nodes": count("nodes"),
            "edges": count("edges"),
            "chunks": count("chunks"),
            "relations": {row["relation"]: row["count"] for row in relation_rows},
            "node_types": {row["node_type"]: row["count"] for row in node_types_rows},
            "db_path": str(self.db_path),
        }

    def retrieve(self, query: str, top_k: int = 10) -> list[dict[str, Any]]:
        query = normalize_space(query)
        if not query:
            return []
        combined: dict[str, float] = {}
        reasons: dict[str, list[str]] = {}

        # A wide candidate pool costs almost nothing now that vector scoring is
        # a single matmul, and it gives the graph stages far more to work with
        # on synthesis questions that need many articles at once.
        pool = max(60, top_k * 8)
        for rank, row in enumerate(self._fts_search(query, limit=pool), start=1):
            score = 1.0 / (rank + 2)
            combined[row["chunk_id"]] = combined.get(row["chunk_id"], 0.0) + score * 1.25
            reasons.setdefault(row["chunk_id"], []).append("FTS")

        try:
            vector_matches = self._vector_search(query, limit=pool)
        except Exception as exc:
            logger.warning(
                "Vertex AI embedding unavailable; skipping local dense retrieval "
                "and continuing with FTS: %s",
                exc,
            )
            vector_matches = []
        for rank, (chunk_id, score) in enumerate(vector_matches, start=1):
            combined[chunk_id] = combined.get(chunk_id, 0.0) + max(score, 0.0) * (1.0 / math.sqrt(rank + 1))
            reasons.setdefault(chunk_id, []).append("vector")

        if not combined:
            return []

        query_ascii = strip_accents(query).lower()
        query_terms = key_terms(query)
        article_numbers = {m.group(1).lower() for m in re.finditer(r"Điều\s+(\d+[a-zA-Z]?)", query, re.I)}
        clause_numbers = {m.group(1) for m in re.finditer(r"khoản\s+(\d{1,3})", query, re.I)}

        doc_priors = self._document_priors(query_ascii)
        rows_by_id = self._chunks_by_ids(combined.keys())
        for chunk_id, row in rows_by_id.items():
            haystack_raw = f"{row['title']} {row['citation']} {row['text'][:600]}"
            haystack = haystack_raw.lower()
            haystack_ascii = strip_accents(haystack_raw).lower()
            if query_terms:
                matched = sum(1 for term in query_terms if term in haystack_ascii)
                coverage = matched / min(len(query_terms), 10)
                combined[chunk_id] += coverage * 0.9
                if coverage < 0.18:
                    combined[chunk_id] *= 0.45
            if "duoc" in query_ascii and "khong duoc" not in query_ascii and "khong duoc" in haystack_ascii:
                combined[chunk_id] -= 0.35
            if "khong duoc" in query_ascii and "khong duoc" in haystack_ascii:
                combined[chunk_id] += 0.5
            if (
                "nguoi su dung lao dong" in query_ascii
                and "don phuong" in query_ascii
                and "cham dut" in query_ascii
                and "quyen don phuong cham dut hop dong lao dong cua nguoi su dung lao dong" in haystack_ascii
            ):
                combined[chunk_id] += 1.15
            if any(term in query_ascii for term in ["lao dong", "hop dong", "nguoi lao dong", "nguoi su dung"]):
                if "lao dong" in haystack_ascii or "hop dong lao dong" in haystack_ascii:
                    combined[chunk_id] += 0.18
            prior = doc_priors.get(str(row["doc_id"]), 1.0)
            if prior != 1.0:
                combined[chunk_id] *= prior
            if row["chunk_type"] in {"article", "clause", "point"}:
                combined[chunk_id] += 0.08
            if row["chunk_type"] == "table":
                # Tables carry the payload numbers (wage brackets, fee scales)
                # that a paraphrased article never repeats.
                combined[chunk_id] += 0.22
            if row["chunk_type"] == "semantic":
                combined[chunk_id] -= 0.10
            if any(token in query_ascii for token in ("bao nhieu", "muc", "toi thieu", "phan tram", "%")):
                if re.search(r"\d", row["text"] or ""):
                    combined[chunk_id] += 0.14
            for number in article_numbers:
                if re.search(rf"điều\s+{re.escape(number)}\b", haystack, re.I):
                    combined[chunk_id] += 0.75
                    reasons.setdefault(chunk_id, []).append("exact-article")
            for number in clause_numbers:
                if re.search(rf"khoản\s+{re.escape(number)}\b", haystack, re.I):
                    combined[chunk_id] += 0.45
                    reasons.setdefault(chunk_id, []).append("exact-clause")

        best_base = max(combined.values(), default=0.0)
        self._seed_definitions(query, query_ascii, combined, reasons, best_base)
        self._seed_concepts(query_ascii, query_terms, combined, reasons, best_base)
        aggregative = self._is_aggregative(query_ascii)

        base_ids = [cid for cid, _ in sorted(combined.items(), key=lambda x: x[1], reverse=True)[: max(top_k * 2, 12)]]
        expanded = self._expand_graph(base_ids, combined, reasons, query_terms, aggregative)
        self._bridge_documents(query, base_ids, expanded, best_base)
        selected = self._diversify(expanded, top_k, aggregative)
        for idx, row in enumerate(selected, start=1):
            row["source_id"] = f"S{idx}"
            row["score"] = round(float(row["score"]), 4)
        return selected

    #: Vietnamese legal hierarchy. A code article states the rule; the decree
    #: only implements it, so when a question does not name the decree the
    #: primary legislation should be cited first. Without this the single
    #: largest implementing decree wins every lexical and vector contest purely
    #: because it is long and repeats the same vocabulary.
    DOC_TYPE_PRIOR = {
        "Bộ luật": 1.16,
        "Luật": 1.12,
        "Văn bản hợp nhất": 1.04,
        "Nghị quyết": 1.0,
        "Nghị định": 0.94,
        "Thông tư": 0.90,
        "Hệ thống": 0.90,
    }

    def _document_priors(self, query_ascii: str) -> dict[str, float]:
        """Per-document score multipliers, neutralised for explicitly named documents."""

        if self._doc_types is None:
            rows = self.conn.execute("SELECT doc_id, doc_type, code FROM docs").fetchall()
            self._doc_types = {
                str(row["doc_id"]): (str(row["doc_type"] or ""), strip_accents(str(row["code"] or "")).lower())
                for row in rows
            }
        priors: dict[str, float] = {}
        for doc_id, (doc_type, code) in self._doc_types.items():
            if code and len(code) >= 6 and code in query_ascii:
                priors[doc_id] = 1.25  # the user asked for this document by name
                continue
            priors[doc_id] = self.DOC_TYPE_PRIOR.get(doc_type, 1.0)
        return priors

    #: Queries that ask for a synthesis rather than a single rule.
    AGGREGATIVE_MARKERS = (
        "tong hop", "liet ke", "day du", "toan bo", "tat ca", "so sanh", "phan biet",
        "nhung khoan", "cac khoan", "nhung nghia vu", "cac nghia vu", "nhung quyen",
        "gom nhung gi", "bao gom nhung", "cac buoc", "nhung truong hop", "cac che do",
        "nhung han h vi", "cac hanh vi", "moc thoi hieu", "ho so rui ro",
    )

    def _is_aggregative(self, query_ascii: str) -> bool:
        return any(marker in query_ascii for marker in self.AGGREGATIVE_MARKERS)

    def _diversify(
        self,
        expanded: dict[str, dict[str, Any]],
        top_k: int,
        aggregative: bool = False,
    ) -> list[dict[str, Any]]:
        """Fill the answer window without letting one document own every slot.

        Vietnamese legal corpora contain a few enormous implementing decrees;
        on a purely score-ordered list they crowd out the very code article the
        question is about. Capping per document forces the cross-document
        evidence a legal answer needs — except for synthesis questions, where
        the answer legitimately lives in many articles of the same code.
        """

        # A concept node's own chunk is a navigation stub ("Thang lương, bảng
        # lương" plus a citation list). It is what leads the walk to the real
        # articles, but it is not evidence, so it must not consume the slots
        # those articles need.
        for row in expanded.values():
            if str(row.get("chunk_type")) == "semantic":
                row["score"] *= 0.5

        ranked = sorted(expanded.values(), key=lambda row: row["score"], reverse=True)
        if aggregative:
            doc_cap = max(4, int(top_k * 0.7))
            node_cap = 1
        else:
            doc_cap = max(3, int(top_k * 0.45))
            node_cap = 2
        semantic_cap = 1
        selected: list[dict[str, Any]] = []
        per_doc: dict[str, int] = {}
        per_node: dict[str, int] = {}
        semantic_used = 0
        deferred: list[dict[str, Any]] = []

        for row in ranked:
            doc_id = str(row.get("doc_id") or "")
            node_id = str(row.get("node_id") or "")
            is_semantic = str(row.get("chunk_type")) == "semantic"
            if is_semantic and semantic_used >= semantic_cap:
                deferred.append(row)
                continue
            if per_doc.get(doc_id, 0) >= doc_cap or per_node.get(node_id, 0) >= node_cap:
                deferred.append(row)
                continue
            selected.append(row)
            per_doc[doc_id] = per_doc.get(doc_id, 0) + 1
            per_node[node_id] = per_node.get(node_id, 0) + 1
            semantic_used += int(is_semantic)
            if len(selected) >= top_k:
                return selected

        # Not enough distinct documents to fill the window - relax the caps.
        for row in deferred:
            selected.append(row)
            if len(selected) >= top_k:
                break
        return selected

    def _load_terms(self) -> list[tuple[str, str]]:
        if self._terms is None:
            rows = self.conn.execute(
                "SELECT node_id, label FROM nodes WHERE node_type = 'ThuậtNgữ'"
            ).fetchall()
            terms = [(strip_accents(row["label"]).lower(), row["node_id"]) for row in rows]
            self._terms = sorted(terms, key=lambda item: -len(item[0]))
        return self._terms

    def _seed_definitions(
        self,
        query: str,
        query_ascii: str,
        combined: dict[str, float],
        reasons: dict[str, list[str]],
        best_base: float,
    ) -> None:
        """Answer "X là gì?" from the mined definition layer, not from keywords.

        Layer 2 knows which clause officially defines each term, so a definition
        question can jump straight there instead of hoping the glossary clause
        happens to out-rank the hundred articles that merely use the term.
        """

        markers = ("la gi", "dinh nghia", "duoc hieu la", "duoc hieu nhu the nao", "khai niem")
        if not any(marker in query_ascii for marker in markers):
            return
        matches = [node_id for term, node_id in self._load_terms() if len(term) >= 6 and term in query_ascii]
        if not matches:
            return
        placeholders = ",".join("?" for _ in matches[:5])
        rows = self.conn.execute(
            f"""
            SELECT c.* FROM edges e
            JOIN chunks c ON c.node_id = e.target_id
            WHERE e.source_id IN ({placeholders}) AND e.relation = 'ĐƯỢC_ĐỊNH_NGHĨA_LÀ'
            """,
            matches[:5],
        ).fetchall()
        boost = max(best_base, 1.0) * 1.15
        for row in rows:
            chunk_id = row["chunk_id"]
            # A clause is the precise definition; the whole glossary article is
            # useful context but should not outrank it.
            weight = 1.0 if row["chunk_type"] == "clause" else 0.72
            combined[chunk_id] = max(combined.get(chunk_id, 0.0), boost * weight)
            reasons.setdefault(chunk_id, []).append("definition")

    #: Semantic node types worth matching against the raw question text.
    SEEDABLE_CONCEPT_TYPES = (
        "KhoảnThuNhập", "LoạiThưởng", "HìnhThứcTrảLương", "KỳHạnTrảLương", "CănCứTínhLương",
        "MứcLươngTốiThiểu", "CáchTính_CôngThức", "ChếĐộ_QuyềnLợi", "NghĩaVụ", "HợpĐồngLaoĐộng",
        "HànhVi_SựKiện", "ThủTục_ChếĐộ", "HồSơ_GiấyTờ", "ĐiềuKiện", "CơQuanGiảiQuyết",
        "HànhViViPhạm", "BiệnPhápKhắcPhục", "TrạngTháiPhápLý", "SựKiệnKíchHoạt",
    )

    def _load_concepts(self) -> list[tuple[str, str]]:
        if self._concepts is None:
            placeholders = ",".join("?" for _ in self.SEEDABLE_CONCEPT_TYPES)
            rows = self.conn.execute(
                f"SELECT node_id, label FROM nodes WHERE node_type IN ({placeholders})",
                list(self.SEEDABLE_CONCEPT_TYPES),
            ).fetchall()
            concepts = [
                (strip_accents(str(row["label"])).lower(), str(row["node_id"]))
                for row in rows
                if len(str(row["label"])) >= 8
            ]
            self._concepts = sorted(concepts, key=lambda item: -len(item[0]))
        return self._concepts

    def _seed_concepts(
        self,
        query_ascii: str,
        query_terms: list[str],
        combined: dict[str, float],
        reasons: dict[str, list[str]],
        best_base: float,
    ) -> None:
        """Pull in the articles the ontology says govern a phrase in the question.

        When a user writes "quy chế thưởng" the graph already records exactly
        which articles regulate it — Điều 104 of the Labour Code, Điều 17 of the
        sanction decree, Điều 41 of the implementing decree. Seeding from the
        concept is far more reliable than hoping all three win a keyword contest
        against a longer, wordier document.
        """

        matched = [node_id for label, node_id in self._load_concepts() if label in query_ascii]
        if not matched:
            return
        boost = max(best_base, 1.0)
        for hub_id in matched[:6]:
            targets = self.conn.execute(
                """
                SELECT target_id FROM edges
                WHERE source_id = ? AND relation IN ('QUY_ĐỊNH_TẠI', 'BỊ_XỬ_PHẠT', 'ÁP_DỤNG_CHO')
                """,
                (hub_id,),
            ).fetchall()
            node_ids = [
                str(row["target_id"])
                for row in targets
                if str(row["target_id"]).startswith(("dieu:", "khoan:"))
            ]
            if not node_ids or len(node_ids) > self.HUB_FANOUT_LIMIT:
                continue
            scored: list[tuple[float, sqlite3.Row]] = []
            for chunk in self._best_chunks_for_nodes(list(dict.fromkeys(node_ids))).values():
                haystack = strip_accents(f"{chunk['title']} {chunk['text'][:500]}").lower()
                coverage = sum(1 for term in query_terms if term in haystack) / min(len(query_terms), 10)
                scored.append((coverage, chunk))
            scored.sort(key=lambda item: -item[0])
            for coverage, chunk in scored[:4]:
                chunk_id = chunk["chunk_id"]
                score = boost * (0.55 + 0.35 * coverage)
                if combined.get(chunk_id, 0.0) < score:
                    combined[chunk_id] = score
                    reasons.setdefault(chunk_id, []).append("concept")

    def _load_doc_bridges(self) -> dict[str, list[str]]:
        """doc_id -> documents linked by GUIDES / AMENDS / REPLACES, both ways."""

        if self._doc_bridges is None:
            bridges: dict[str, list[str]] = {}
            rows = self.conn.execute(
                "SELECT source_id, target_id FROM edges WHERE relation IN ('HƯỚNG_DẪN', 'SỬA_ĐỔI', 'THAY_THẾ')"
            ).fetchall()
            for row in rows:
                source = str(row["source_id"]).removeprefix("doc:")
                target = str(row["target_id"]).removeprefix("doc:")
                bridges.setdefault(source, []).append(target)
                bridges.setdefault(target, []).append(source)
            self._doc_bridges = {key: list(dict.fromkeys(value)) for key, value in bridges.items()}
        return self._doc_bridges

    def _bridge_documents(
        self,
        query: str,
        base_ids: list[str],
        expanded: OrderedDict[str, dict[str, Any]],
        best_base: float,
    ) -> None:
        """Re-query the linked law when the hits all sit in one decree.

        A question like "trả lương chậm bị phạt bao nhiêu" is answered by the
        sanction decree *and* the code article it enforces. Following the
        document-level GUIDES edge and running a second, document-scoped search
        pulls the missing half in at article granularity.
        """

        expression = self._fts_query(query)
        if not expression:
            return
        base_rows = self._chunks_by_ids(base_ids)
        present = {str(row["doc_id"]) for row in base_rows.values()}
        bridges = self._load_doc_bridges()
        candidates: list[str] = []
        for doc_id in present:
            for linked in bridges.get(doc_id, []):
                if linked not in present and linked not in candidates:
                    candidates.append(linked)
        if not candidates:
            return

        for doc_id in candidates[:4]:
            try:
                rows = self.conn.execute(
                    """
                    SELECT c.*, bm25(chunk_fts) AS rank
                    FROM chunk_fts
                    JOIN chunks c ON c.chunk_id = chunk_fts.chunk_id
                    WHERE chunk_fts MATCH ? AND c.doc_id = ? AND c.chunk_type IN ('article', 'clause', 'table')
                    ORDER BY rank
                    LIMIT 2
                    """,
                    (expression, doc_id),
                ).fetchall()
            except sqlite3.OperationalError:
                continue
            for position, row in enumerate(rows):
                chunk_id = row["chunk_id"]
                score = best_base * (0.62 - position * 0.06)
                current = expanded.get(chunk_id)
                if current and current["score"] >= score:
                    continue
                payload = dict(row)
                payload.pop("rank", None)
                payload["score"] = score
                payload["reasons"] = ["bridge:HƯỚNG_DẪN"]
                expanded[chunk_id] = payload

    def _fts_query(self, query: str) -> str:
        tokens = [t for t in VN_WORD_RE.findall(query.lower()) if len(t) >= 2]
        stop = {
            "theo",
            "quy",
            "dinh",
            "quy định",
            "cho",
            "toi",
            "tôi",
            "hoi",
            "hỏi",
            "nhu",
            "như",
            "nao",
            "nào",
            "ve",
            "về",
        }
        cleaned = []
        for token in tokens:
            if strip_accents(token) in stop:
                continue
            cleaned.append(token)
        cleaned = cleaned[:18] or tokens[:12]
        return " OR ".join(f'"{token}"' for token in cleaned)

    def _fts_search(self, query: str, limit: int) -> list[sqlite3.Row]:
        expr = self._fts_query(query)
        if not expr:
            return []
        try:
            return self.conn.execute(
                """
                SELECT c.*, bm25(chunk_fts) AS rank
                FROM chunk_fts
                JOIN chunks c ON c.chunk_id = chunk_fts.chunk_id
                WHERE chunk_fts MATCH ?
                ORDER BY rank
                LIMIT ?
                """,
                (expr, limit),
            ).fetchall()
        except sqlite3.OperationalError:
            like = f"%{query[:80]}%"
            return self.conn.execute(
                "SELECT * FROM chunks WHERE text LIKE ? OR title LIKE ? LIMIT ?",
                (like, like, limit),
            ).fetchall()

    def _load_vectors(self) -> list[tuple[str, bytes]]:
        if self._vectors is None:
            rows = self.conn.execute("SELECT chunk_id, vector FROM chunks").fetchall()
            self._vectors = [(row["chunk_id"], row["vector"]) for row in rows]
        return self._vectors

    def _load_matrix(self) -> Any:
        """Stack every chunk vector once so a query is a single matmul.

        Scoring 27k chunks with a Python loop cost ~2.5s per query; as one numpy
        product it is a few milliseconds, which is the difference between a
        usable interactive latency and a visible stall.
        """

        if self._matrix is not None:
            return self._matrix
        try:
            import numpy as np
        except ImportError:  # pragma: no cover - numpy is a hard dependency of torch
            self._matrix = False
            return self._matrix

        rows = self._load_vectors()
        dimensions = self.embedding_config.dimensions
        matrix = np.empty((len(rows), dimensions), dtype=np.float32)
        for index, (chunk_id, blob) in enumerate(rows):
            vector = np.frombuffer(blob, dtype=np.float32)
            if vector.size != dimensions:
                raise RuntimeError(
                    f"Chunk {chunk_id} has {vector.size} embedding dimensions; rebuild the index."
                )
            matrix[index] = vector
        self._matrix = matrix
        self._matrix_ids = [chunk_id for chunk_id, _ in rows]
        return self._matrix

    def _vector_search(self, query: str, limit: int) -> list[tuple[str, float]]:
        qvec = get_embedding_service(self.embedding_config).embed_query(query)
        matrix = self._load_matrix()
        if matrix is not False:
            import numpy as np

            scores = matrix @ np.asarray(qvec, dtype=np.float32)
            count = min(limit, scores.shape[0])
            if count <= 0:
                return []
            top = np.argpartition(-scores, count - 1)[:count]
            top = top[np.argsort(-scores[top])]
            return [(self._matrix_ids[index], float(scores[index])) for index in top]

        scored = []
        for chunk_id, blob in self._load_vectors():
            vector = blob_to_vector(blob)
            if len(vector) != self.embedding_config.dimensions:
                raise RuntimeError(
                    f"Chunk {chunk_id} has {len(vector)} embedding dimensions; rebuild the index."
                )
            scored.append((chunk_id, dot(qvec, vector)))
        scored.sort(key=lambda x: x[1], reverse=True)
        return scored[:limit]

    def _chunks_by_ids(self, chunk_ids: Any) -> dict[str, sqlite3.Row]:
        ids = list(dict.fromkeys(chunk_ids))
        if not ids:
            return {}
        placeholders = ",".join("?" for _ in ids)
        rows = self.conn.execute(f"SELECT * FROM chunks WHERE chunk_id IN ({placeholders})", ids).fetchall()
        return {row["chunk_id"]: row for row in rows}

    CHUNK_PRIORITY = ["point", "clause", "article", "table", "sliding", "document_intro", "structure", "semantic"]

    def _rank_chunk(self, row: sqlite3.Row) -> tuple[int, int]:
        chunk_type = row["chunk_type"]
        order = self.CHUNK_PRIORITY.index(chunk_type) if chunk_type in self.CHUNK_PRIORITY else 99
        return order, row["ordinal"]

    def _best_chunk_for_node(self, node_id: str) -> sqlite3.Row | None:
        """Representative chunk for a node, memoised per store instance.

        Graph expansion asks for this thousands of times per query; without the
        cache each ask was a separate SQLite round trip.
        """

        if node_id in self._node_chunk_cache:
            return self._node_chunk_cache[node_id]
        rows = self.conn.execute("SELECT * FROM chunks WHERE node_id = ?", (node_id,)).fetchall()
        best = min(rows, key=self._rank_chunk) if rows else None
        self._node_chunk_cache[node_id] = best
        return best

    def _best_chunks_for_nodes(self, node_ids: list[str]) -> dict[str, sqlite3.Row]:
        """Batched form of :meth:`_best_chunk_for_node` for wide fan-outs."""

        missing = [node_id for node_id in node_ids if node_id not in self._node_chunk_cache]
        for start in range(0, len(missing), 400):
            window = missing[start : start + 400]
            placeholders = ",".join("?" for _ in window)
            rows = self.conn.execute(
                f"SELECT * FROM chunks WHERE node_id IN ({placeholders})", window
            ).fetchall()
            grouped: dict[str, sqlite3.Row] = {}
            for row in rows:
                node_id = row["node_id"]
                current = grouped.get(node_id)
                if current is None or self._rank_chunk(row) < self._rank_chunk(current):
                    grouped[node_id] = row
            for node_id in window:
                self._node_chunk_cache[node_id] = grouped.get(node_id)
        return {
            node_id: self._node_chunk_cache[node_id]
            for node_id in node_ids
            if self._node_chunk_cache.get(node_id) is not None
        }

    def _ancestor_nodes(self, node_id: str) -> list[str]:
        ancestors: list[str] = []
        cursor = node_id
        seen: set[str] = set()
        while cursor and cursor not in seen:
            seen.add(cursor)
            row = self.conn.execute("SELECT parent_id FROM nodes WHERE node_id = ?", (cursor,)).fetchone()
            if not row or not row["parent_id"]:
                break
            cursor = row["parent_id"]
            ancestors.append(cursor)
        return ancestors

    def _node_edges(self, node_ids: list[str]) -> list[sqlite3.Row]:
        if not node_ids:
            return []
        placeholders = ",".join("?" for _ in node_ids)
        reverse_rels = sorted(onto.REVERSIBLE_RELATIONS)
        reverse_placeholders = ",".join("?" for _ in reverse_rels)
        return self.conn.execute(
            f"""
            SELECT * FROM edges
            WHERE source_id IN ({placeholders})
               OR (target_id IN ({placeholders}) AND relation IN ({reverse_placeholders}))
            """,
            node_ids + node_ids + reverse_rels,
        ).fetchall()

    def _sibling_articles(self, node_id: str, limit: int = 2) -> list[sqlite3.Row]:
        """Articles immediately before/after this one under the same parent.

        Related rules are written side by side - trợ cấp thôi việc (Điều 46) and
        trợ cấp mất việc làm (Điều 47), điều kiện hưởng lương hưu and mức lương
        hưu - so a question that contrasts two of them needs both.
        """

        row = self.conn.execute(
            "SELECT parent_id, ordinal, node_type FROM nodes WHERE node_id = ?", (node_id,)
        ).fetchone()
        if not row or row["node_type"] != "Điều" or not row["parent_id"]:
            return []
        return self.conn.execute(
            """
            SELECT node_id, label, ordinal FROM nodes
            WHERE parent_id = ? AND node_type = 'Điều' AND node_id != ?
            ORDER BY ABS(ordinal - ?) LIMIT ?
            """,
            (row["parent_id"], node_id, row["ordinal"], limit),
        ).fetchall()

    #: A concept linked to more articles than this is a generic hub ("người lao
    #: động" touches 2.7k articles) and says nothing about relevance.
    HUB_FANOUT_LIMIT = 60

    def _expand_through_hubs(
        self,
        node_ids: list[str],
        base_score: float,
        query_terms: list[str],
        add: Any,
        aggregative: bool = False,
    ) -> None:
        """Two-hop jump: article → shared concept → the other articles about it.

        This is what lets "quy chế thưởng" connect Điều 104 of the Labour Code
        to Điều 17 of the sanction decree and Điều 41 of the implementing
        decree, none of which cite each other textually.

        For an aggregative question ("tổng hợp…", "so sánh…") the broad topic
        hubs are exactly the right structure — they enumerate every article on a
        subject — so the fan-out guard is relaxed and more targets are taken.
        """

        placeholders = ",".join("?" for _ in node_ids)
        hub_relations = ("QUY_ĐỊNH_TẠI", "BỊ_XỬ_PHẠT", "ĐƯỢC_ĐỊNH_NGHĨA_LÀ", "ÁP_DỤNG_CHO", "THUỘC_CHỦ_ĐỀ")
        rel_placeholders = ",".join("?" for _ in hub_relations)
        hubs = self.conn.execute(
            f"""
            SELECT DISTINCT source_id AS hub FROM edges
            WHERE target_id IN ({placeholders}) AND relation IN ({rel_placeholders})
            UNION
            SELECT DISTINCT target_id AS hub FROM edges
            WHERE source_id IN ({placeholders}) AND relation = 'THUỘC_CHỦ_ĐỀ'
            """,
            list(node_ids) + list(hub_relations) + list(node_ids),
        ).fetchall()
        hub_ids = list(dict.fromkeys(str(row["hub"]) for row in hubs))[:10]
        if not hub_ids:
            return

        fanout_limit = 4000 if aggregative else self.HUB_FANOUT_LIMIT
        take = 5 if aggregative else 2
        seen = set(node_ids)
        for hub_id in hub_ids:
            is_topic = hub_id.startswith("chude:")
            if is_topic and not aggregative:
                # Topic hubs are too broad to help a specific question.
                continue
            targets = self.conn.execute(
                """
                SELECT source_id, target_id, relation FROM edges
                WHERE (source_id = ? AND relation IN ('QUY_ĐỊNH_TẠI', 'BỊ_XỬ_PHẠT', 'ĐƯỢC_ĐỊNH_NGHĨA_LÀ'))
                   OR (target_id = ? AND relation = 'THUỘC_CHỦ_ĐỀ')
                """,
                (hub_id, hub_id),
            ).fetchall()
            if not targets or len(targets) > fanout_limit:
                continue
            candidates = [
                str(row["target_id"] if row["relation"] != "THUỘC_CHỦ_ĐỀ" else row["source_id"])
                for row in targets
            ]
            candidates = [
                node_id
                for node_id in dict.fromkeys(candidates)
                if node_id not in seen and node_id.startswith(("dieu:", "khoan:"))
            ][:600]
            scored: list[tuple[float, sqlite3.Row]] = []
            for chunk in self._best_chunks_for_nodes(candidates).values():
                haystack = strip_accents(f"{chunk['title']} {chunk['text'][:500]}").lower()
                coverage = sum(1 for term in query_terms if term in haystack) / min(len(query_terms), 10)
                if coverage < (0.2 if aggregative else 0.25):
                    continue
                scored.append((coverage, chunk))
            scored.sort(key=lambda item: -item[0])
            weight_base = 0.50 if aggregative else 0.44
            for coverage, chunk in scored[:take]:
                add(chunk, base_score * (weight_base + 0.18 * coverage), f"hub:{hub_id.split(':')[0]}")

    def _expand_graph(
        self,
        base_ids: list[str],
        base_scores: dict[str, float],
        reasons: dict[str, list[str]],
        query_terms: list[str] | None = None,
        aggregative: bool = False,
    ) -> OrderedDict[str, dict[str, Any]]:
        expanded: OrderedDict[str, dict[str, Any]] = OrderedDict()
        base_rows = self._chunks_by_ids(base_ids)

        def add(row: sqlite3.Row, score: float, reason: str) -> None:
            current = expanded.get(row["chunk_id"])
            if current and current["score"] >= score:
                if reason not in current["reasons"]:
                    current["reasons"].append(reason)
                return
            payload = dict(row)
            payload["score"] = score
            payload["reasons"] = list(dict.fromkeys(reasons.get(row["chunk_id"], []) + [reason]))
            expanded[row["chunk_id"]] = payload

        for chunk_id in base_ids:
            row = base_rows.get(chunk_id)
            if not row:
                continue
            base_score = base_scores.get(chunk_id, 0.0)
            add(row, base_score, "base")
            node_ids = [row["node_id"]] + self._ancestor_nodes(row["node_id"])
            for pos, ancestor_id in enumerate(node_ids[1:], start=1):
                ancestor_chunk = self._best_chunk_for_node(ancestor_id)
                if ancestor_chunk:
                    if ancestor_chunk["chunk_type"] == "article":
                        weight = 1.05
                    elif ancestor_chunk["chunk_type"] == "clause":
                        weight = 0.9
                    elif ancestor_chunk["chunk_type"] == "structure":
                        weight = 0.38
                    else:
                        weight = max(0.28, 0.72 - pos * 0.1)
                    add(ancestor_chunk, base_score * weight, "ancestor")
            for edge in self._node_edges(node_ids):
                reverse = edge["source_id"] not in node_ids
                other_id = edge["source_id"] if reverse else edge["target_id"]
                if edge["relation"] == "BAN_HÀNH":
                    continue
                edge_chunk = self._best_chunk_for_node(other_id)
                if edge_chunk:
                    weight = onto.relation_weight(edge["relation"])
                    if reverse:
                        # Walking a relation backwards is a weaker signal than
                        # following it in its declared direction.
                        weight *= 0.85
                    add(edge_chunk, base_score * weight, f"edge:{edge['relation']}")

            if query_terms:
                self._expand_through_hubs(node_ids, base_score, query_terms, add, aggregative)

            article_id = next(
                (nid for nid in node_ids if str(nid).startswith("dieu:")), None
            )
            if article_id and query_terms:
                for sibling in self._sibling_articles(article_id, limit=4):
                    sibling_chunk = self._best_chunk_for_node(sibling["node_id"])
                    if not sibling_chunk:
                        continue
                    haystack = strip_accents(
                        f"{sibling_chunk['title']} {sibling_chunk['text'][:400]}"
                    ).lower()
                    matched = sum(1 for term in query_terms if term in haystack)
                    if matched / min(len(query_terms), 10) < 0.3:
                        continue
                    add(sibling_chunk, base_score * 0.58, "sibling")

        return expanded


def build_index(
    data_dir: Path | str | None = None,
    storage_dir: Path | str | None = None,
) -> dict[str, int]:
    builder = LegalGraphBuilder(
        Path(data_dir or os.getenv("LEGAL_DATA_DIR", DEFAULT_DATA_DIR)),
        Path(storage_dir or os.getenv("LEGAL_STORAGE_DIR", DEFAULT_STORAGE_DIR)),
        EmbeddingConfig.from_env(),
    )
    return builder.build()
