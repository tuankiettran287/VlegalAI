from __future__ import annotations

import re
import sys
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "BAO_CAO_DANH_GIA_VLEGALAI_RAG_GRAPHRAG_V2.md"
OUTPUT = ROOT / "BAO_CAO_DANH_GIA_VLEGALAI_RAG_GRAPHRAG_V2.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
GRAY = "667085"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8F1FB"
LIGHT_INDIGO = "EEF2FF"
GREEN = "E7F6EC"
AMBER = "FFF4D6"
RED = "FDEBEC"
NA_GRAY = "EAECF0"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float], table_width_twips: int) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    normalized = [w / sum(widths) for w in widths]
    twips = [int(table_width_twips * w) for w in normalized]
    twips[-1] += table_width_twips - sum(twips)

    grid_cols = list(table._tbl.tblGrid.gridCol_lst)
    for col, width in zip(grid_cols, twips):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, twips[idx])
            cell.width = Inches(twips[idx] / 1440)


def set_table_borders(table, color: str = "D0D5DD", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_header_and_keep(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for child in (begin, instr, separate, text, end):
        run._r.append(child)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Mục lục sẽ được cập nhật khi mở tài liệu."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for child in (begin, instr, separate, placeholder, end):
        run._r.append(child)


def set_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def configure_section(section, mode: str) -> None:
    if mode == "PORTRAIT":
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
    elif mode == "LANDSCAPE":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)
    elif mode == "TABLOID":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(17)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.header_distance = Inches(0.28)
        section.footer_distance = Inches(0.28)
    else:
        raise ValueError(mode)


def add_running_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.clear()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left = header_para.add_run("VLEGALAI")
    left.bold = True
    left.font.name = "Calibri"
    left.font.size = Pt(8.5)
    left.font.color.rgb = RGBColor.from_string(BLUE)
    middle = header_para.add_run("    ARCHITECTURE & AI ENGINEERING REVIEW")
    middle.font.name = "Calibri"
    middle.font.size = Pt(8)
    middle.font.color.rgb = RGBColor.from_string(GRAY)
    right = header_para.add_run("    REVISION 2.1")
    right.bold = True
    right.font.name = "Calibri"
    right.font.size = Pt(8)
    right.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    header_para.paragraph_format.space_after = Pt(0)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = footer_para.add_run("Internal readiness report  •  ")
    prefix.font.name = "Calibri"
    prefix.font.size = Pt(8)
    prefix.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_field(footer_para)


def add_section(doc: Document, mode: str):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, mode)
    add_running_header_footer(section)
    return section


def available_width_twips(doc: Document) -> int:
    section = doc.sections[-1]
    width_emu = section.page_width - section.left_margin - section.right_margin
    return int(width_emu / 635)


def add_inline(paragraph, text: str, *, size: float | None = None, color: str | None = None) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\[[^\]]+\]\([^\)]+\))")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt((size or 10) - 0.5)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token)
            run = paragraph.add_run(link.group(1) if link else token)
            run.underline = True
            run.font.color.rgb = RGBColor.from_string(BLUE)
        if size and not token.startswith("`"):
            run.font.size = Pt(size)
        if color and not token.startswith("`"):
            run.font.color.rgb = RGBColor.from_string(color)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def table_width_profile(headers: list[str]) -> list[float]:
    n = len(headers)
    normalized = [re.sub(r"[`*]", "", h).strip().lower() for h in headers]
    if normalized == ["id", "mandatory", "trạng thái", "đánh giá"]:
        return [0.09, 0.10, 0.19, 0.62]
    if n == 8 and "endpoint" in normalized[0]:
        return [0.13, 0.11, 0.13, 0.14, 0.10, 0.10, 0.09, 0.20]
    if n == 8 and any("latency" in h for h in normalized):
        return [0.22, 0.08, 0.08, 0.12, 0.12, 0.12, 0.12, 0.14]
    if n == 6 and normalized[0] == "phạm vi":
        return [0.32, 0.13, 0.16, 0.16, 0.15, 0.08]
    if n == 5 and normalized[0] in {"entity", "nhóm"}:
        return [0.17, 0.14, 0.22, 0.27, 0.20]
    if n == 5 and normalized[0].startswith("store"):
        return [0.17, 0.22, 0.16, 0.21, 0.24]
    if n == 5 and normalized[0] == "ưu tiên":
        return [0.08, 0.24, 0.16, 0.52]
    if n == 5:
        return [0.18, 0.16, 0.18, 0.26, 0.22]
    if n == 4:
        return [0.18, 0.22, 0.22, 0.38]
    if n == 3:
        return [0.23, 0.35, 0.42]
    if n == 2:
        return [0.32, 0.68]
    return [1 / n] * n


def status_fill(text: str) -> str | None:
    plain = re.sub(r"[`*]", "", text)
    if "Chưa đạt" in plain:
        return RED
    if "Một phần" in plain:
        return AMBER
    if "Không áp dụng" in plain:
        return NA_GRAY
    if "Đạt" in plain:
        return GREEN
    return None


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    headers, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    widths = table_width_profile(headers)
    set_table_geometry(table, widths, available_width_twips(doc))
    font_size = 7.5 if len(headers) >= 7 else 8.0 if len(headers) >= 5 else 8.5 if len(headers) == 4 else 9.0

    header_row = table.rows[0]
    repeat_table_header(header_row)
    keep_row_together(header_row)
    for idx, value in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        add_inline(p, value, size=font_size, color=DARK_BLUE)
        for run in p.runs:
            run.bold = True

    for values in body:
        row = table.add_row()
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            fill = status_fill(value)
            if fill and ("trạng thái" in headers[idx].lower() or idx == 2):
                set_cell_shading(cell, fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.02
            add_inline(p, value, size=font_size)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)
    after.paragraph_format.space_before = Pt(0)


def add_diagram(doc: Document, image_spec: str) -> None:
    payload = image_spec[len("[[DIAGRAM:") : -2]
    rel_path, caption, requested_width = payload.split("|", 2)
    image_path = ROOT / rel_path
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    requested = float(requested_width)
    section = doc.sections[-1]
    max_width = (section.page_width - section.left_margin - section.right_margin) / 914400
    max_height = (section.page_height - section.top_margin - section.bottom_margin) / 914400 - 0.75
    with Image.open(image_path) as image:
        px_w, px_h = image.size
    width = min(requested, max_width)
    height = width * px_h / px_w
    if height > max_height:
        height = max_height
        width = height * px_w / px_h
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shape = p.add_run().add_picture(str(image_path), width=Inches(width), height=Inches(height))
    set_alt_text(shape, caption.split("—", 1)[0].strip(), caption)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    add_inline(cap, caption, size=9, color=GRAY)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(6)

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.italic = True
    callout.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    callout.paragraph_format.left_indent = Inches(0.25)
    callout.paragraph_format.right_indent = Inches(0.25)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)


def add_cover(doc: Document) -> None:
    bar = doc.add_table(rows=1, cols=1)
    bar.autofit = False
    repeat_table_header(bar.rows[0])
    set_table_geometry(bar, [1], available_width_twips(doc))
    set_cell_shading(bar.cell(0, 0), BLUE)
    bar.cell(0, 0).paragraphs[0].add_run(" ")
    bar.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("VLEGALAI")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("BÁO CÁO ĐÁNH GIÁ KIẾN TRÚC,\nTHIẾT KẾ VÀ AI ENGINEERING")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor.from_string(INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("RAG • GraphRAG đa tầng • HybridRAG • DeepSeek 7B")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    meta = doc.add_table(rows=6, cols=2)
    repeat_table_header(meta.rows[0])
    set_table_borders(meta, color="D0D5DD", size="4")
    set_table_geometry(meta, [0.28, 0.72], available_width_twips(doc))
    values = [
        ("Dự án", "VLegalAI / LaborCare GraphRAG"),
        ("Phiên bản", "2.1 — Physical Database Design revision"),
        ("Ngày", "13/07/2026"),
        ("Model mục tiêu", "DeepSeek-R1-Distill-Qwen-7B"),
        ("Product modes", "RAG, GRAPHRAG, HYBRID_RAG"),
        ("Phạm vi", r"F:\VlegalAI"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell in row.cells:
            set_cell_margins(cell, top=100, bottom=100)
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        add_inline(p1, label, size=9.5, color=DARK_BLUE)
        for run in p1.runs:
            run.bold = True
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        add_inline(p2, value, size=9.5)

    doc.add_paragraph()
    note = doc.add_table(rows=1, cols=1)
    repeat_table_header(note.rows[0])
    set_table_geometry(note, [1], available_width_twips(doc))
    set_cell_shading(note.cell(0, 0), LIGHT_BLUE)
    set_cell_margins(note.cell(0, 0), top=180, start=120, bottom=180, end=120)
    p = note.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(
        p,
        "Revision 2.1 bổ sung PostgreSQL Physical Database Design cho user/chat, hash-only PII và conversation, DDL V001 và checklist chứng minh materialization.",
        size=10.5,
        color=DARK_BLUE,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("INTERNAL READINESS REVIEW")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    run.bold = True


def parse_markdown_into_doc(doc: Document, text: str) -> None:
    lines = text.splitlines()
    try:
        index = lines.index("## 1. Kiểm soát tài liệu")
    except ValueError as exc:
        raise RuntimeError("Report source is missing section 1") from exc

    i = index
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "[[PAGE_BREAK]]":
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue
        if stripped.startswith("[[SECTION:"):
            mode = stripped[len("[[SECTION:") : -2]
            add_section(doc, mode)
            i += 1
            continue
        if stripped.startswith("[[DIAGRAM:"):
            add_diagram(doc, stripped)
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:])
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, stripped[3:])
            i += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[5:])
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i + 1]):
            rows = [parse_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, rows)
            continue
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Callout")
            add_inline(p, stripped.lstrip("> "))
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            i += 1
            continue
        match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if match:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, match.group(2))
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1


def apply_document_settings(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    props = doc.core_properties
    props.title = "Báo cáo đánh giá VLegalAI RAG/GraphRAG/HybridRAG — Revision 2.1"
    props.subject = "Architecture, detailed design and AI engineering readiness review"
    props.author = "VLegalAI Project"
    props.last_modified_by = "VLegalAI Project"
    props.keywords = "VLegalAI, RAG, GraphRAG, HybridRAG, DeepSeek 7B, ERD, Mermaid"
    props.comments = "Generated from repository evidence; contains no credential values."


def build() -> Path:
    doc = Document()
    configure_section(doc.sections[0], "PORTRAIT")
    add_running_header_footer(doc.sections[0])
    configure_styles(doc)
    apply_document_settings(doc)
    add_cover(doc)

    doc.add_page_break()
    toc_title = doc.add_paragraph()
    toc_title.paragraph_format.space_after = Pt(12)
    run = toc_title.add_run("MỤC LỤC")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    toc = doc.add_paragraph()
    add_toc(toc)
    doc.add_page_break()

    source_text = SOURCE.read_text(encoding="utf-8")
    parse_markdown_into_doc(doc, source_text)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(path)
